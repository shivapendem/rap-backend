import os
import io
import re
import uuid
import math
import asyncio
from typing import Optional, List
from datetime import datetime, timezone, date
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Body
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import func, or_

from database import get_db
from models import User, Resume, ConsultantExperience, Consultant, RecruiterConsultant
from auth import get_current_user
from s3_service import upload_file_to_s3, generate_presigned_url, delete_file_from_s3, download_file_from_s3
from claude_service import generate_tailored_resume
from phase8_ai_usage_service import save_claude_rate_limits
from resume_validation import get_missing_resume_fields, missing_fields_message
from phase3 import _extract_text_from_docx

# You can import openai and use it if an API key is provided
# import openai

router = APIRouter(prefix="/api/resume", tags=["resume"])

class ResumeCreateRequest(BaseModel):
    title: str
    target_role: Optional[str] = None
    job_description: Optional[str] = None
    experience_ids: Optional[List[int]] = [] # IDs of ConsultantExperience to include
    draft: bool = False # If True, don't generate PDF yet
    user_id: Optional[int] = None # The candidate user_id to generate for
    # Set when generated from the dashboard's "no job description" custom
    # resume flow — links this Resume back to that specific requirement so
    # the dashboard can find it and unlock the Apply button.
    requirement_id: Optional[int] = None

async def _get_resume_for_user(db: AsyncSession, resume_id: int, current_user: User):
    if current_user.role == "ADMIN":
        result = await db.execute(select(Resume).where(Resume.id == resume_id))
        return result.scalar_one_or_none()
    elif current_user.role == "RECRUITER":
        consultant_users_query = select(Consultant.user_id).where(
            or_(
                Consultant.sales_recruiter_user_id == current_user.id,
                Consultant.id.in_(
                    select(RecruiterConsultant.consultant_id).where(
                        RecruiterConsultant.recruiter_id == current_user.id
                    )
                )
            )
        )
        result = await db.execute(select(Resume).where(
            Resume.id == resume_id,
            or_(
                Resume.user_id == current_user.id,
                Resume.user_id.in_(consultant_users_query)
            )
        ))
        return result.scalar_one_or_none()
    else:
        result = await db.execute(select(Resume).where(Resume.id == resume_id, Resume.user_id == current_user.id))
        return result.scalar_one_or_none()

class ResumeUpdateRequest(BaseModel):
    title: Optional[str] = None
    target_role: Optional[str] = None
    job_description: Optional[str] = None
    data: Optional[dict] = None
    status: Optional[str] = None

class ResumeResponse(BaseModel):
    id: int
    user_id: int
    requirement_id: Optional[int] = None
    title: str
    target_role: Optional[str] = None
    job_description: Optional[str] = None
    data: dict
    s3_key: Optional[str] = None
    s3_url: Optional[str] = None
    ats_score: Optional[int] = None
    status: str
    download_count: int
    last_downloaded: Optional[datetime] = None
    created_at: datetime
    updated_at: datetime
    is_base: bool = False

    class Config:
        from_attributes = True

class PaginatedResumes(BaseModel):
    data: List[ResumeResponse]
    total: int
    page: int
    page_size: int
    total_pages: int

async def _resolve_target_user(
    request_user_id: Optional[int],
    current_user: User,
    db: AsyncSession,
):
    """Resolve which consultant we're generating for, honoring the
    ADMIN/RECRUITER "generate on behalf of" path. Shared by the /generate
    endpoint and the /completeness pre-check so they never drift apart.

    BUG FIX: this never checked that a RECRUITER was actually assigned to
    the target consultant — any recruiter could pass any user_id in the
    request/URL and this would happily resolve it. That's how a recruiter
    could reach the "Generate Resume" page and even get as far as the
    completeness check for a consultant who then failed to load on the
    Consultant Detail page ("may not be assigned to you") — two different
    endpoints enforcing two different rules for the same relationship.
    _get_resume_for_user() above already has the correct check (a
    consultant counts as assigned via EITHER the legacy single
    sales_recruiter_user_id FK OR the RecruiterConsultant join table) —
    reuse that same rule here so every recruiter-facing endpoint in this
    file agrees on who's assigned to whom.
    """
    target_user_id = current_user.id
    target_user = current_user
    if request_user_id and current_user.role in ("ADMIN", "RECRUITER"):
        target_user_id = request_user_id
        target_user_result = await db.execute(select(User).where(User.id == target_user_id))
        target_user = target_user_result.scalar_one_or_none()
        if not target_user:
            raise HTTPException(status_code=404, detail="Target user not found")

        if current_user.role == "RECRUITER":
            assigned_consultant = await db.execute(
                select(Consultant.id).where(
                    Consultant.user_id == target_user_id,
                    or_(
                        Consultant.sales_recruiter_user_id == current_user.id,
                        Consultant.id.in_(
                            select(RecruiterConsultant.consultant_id).where(
                                RecruiterConsultant.recruiter_id == current_user.id
                            )
                        ),
                    ),
                )
            )
            if not assigned_consultant.scalar_one_or_none():
                raise HTTPException(status_code=403, detail="This consultant is not assigned to you.")

    if target_user.role != "CONSULTANT":
        raise HTTPException(status_code=403, detail="Resumes can only be generated for consultants.")

    return target_user_id, target_user


async def _build_resume_info(
    target_user: User,
    db: AsyncSession,
    experience_ids: Optional[List[int]] = None,
):
    """Assemble the resume_info payload (skills/experience/education/etc.)
    used both to actually call the AI and to check profile completeness."""
    consultant_res = await db.execute(select(Consultant).where(Consultant.user_id == target_user.id))
    consultant = consultant_res.scalar_one_or_none()

    profile_experiences = []
    if consultant:
        exp_results = await db.execute(select(ConsultantExperience).where(ConsultantExperience.consultant_id == consultant.id).order_by(ConsultantExperience.sort_order))
        profile_experiences = exp_results.scalars().all()

    explicit_experiences = []
    if experience_ids:
        exp_results = await db.execute(select(ConsultantExperience).where(ConsultantExperience.id.in_(experience_ids)))
        explicit_experiences = exp_results.scalars().all()

    # Merge avoiding duplicates
    all_exps = {exp.id: exp for exp in profile_experiences}
    for exp in explicit_experiences:
        all_exps[exp.id] = exp

    merged_exps = list(all_exps.values())

    manual_exp_entries = []
    for exp in merged_exps:
        bullets = []
        if exp.responsibilities:
            bullets.extend([b.strip() for b in exp.responsibilities.split('\n') if b.strip()])
        if exp.achievements:
            bullets.extend([b.strip() for b in exp.achievements.split('\n') if b.strip()])

        tech_str = ", ".join(exp.technologies) if exp.technologies else ""
        if tech_str:
            bullets.append(f"Technologies: {tech_str}")

        start_str = exp.start_date.strftime("%b %Y") if exp.start_date else ""
        end_str = "Present" if exp.is_present else (exp.end_date.strftime("%b %Y") if exp.end_date else "")
        date_str = f"{start_str} - {end_str}".strip(" -")

        # BUG FIX: this used to send ONLY a combined "dates" string (e.g.
        # "Jan 2020 - Present"), but the AI is asked to RETURN separate
        # "start"/"end" fields — forcing it to re-parse and split a
        # combined string on its own instead of just carrying forward
        # values it was already given cleanly. Sending both the combined
        # string (kept for any code that still reads it) AND explicit
        # start/end/is_present makes it far more likely a past role's
        # real end date survives the AI round-trip correctly.
        manual_exp_entries.append({
            "title": exp.role_title,
            "company": exp.client_name,
            "dates": date_str,
            "start": start_str,
            "end": end_str,
            "is_present": bool(exp.is_present),
            "bullets": bullets
        })

    # Use resume_info if available, else build a basic profile from db
    import copy
    resume_info = copy.deepcopy(target_user.resume_info) if target_user.resume_info else {
        "full_name": target_user.full_name or target_user.email.split('@')[0],
        "email": target_user.email,
        "experience": []
    }

    if "experience" not in resume_info:
        resume_info["experience"] = []

    # Prepend profile experiences so they are processed as most relevant/recent
    resume_info["experience"] = manual_exp_entries + resume_info["experience"]

    # BUG FIX: this function fetched `consultant` above but only ever used
    # it for ConsultantExperience rows — it never merged the consultant's
    # own real profile columns (linkedin_url, phone, total_experience_years,
    # primary_skills, preferred_roles) into resume_info at all. Those
    # columns are what the consultant profile screen actually reads/writes
    # (see phase3.py, phase_users_service.py); resume_info is a separate
    # JSON blob that only ever reflects whatever was parsed/imported at
    # signup. So a consultant who filled in LinkedIn (or phone, years of
    # experience, skills, target role) through their profile after that
    # initial import still failed the completeness check below and the
    # AI generation call itself never saw those fields — "Add LinkedIn to
    # the profile" fired even though it genuinely was on file, just not in
    # this blob. Backfill every field this checks against from the real
    # Consultant row, same "real column wins over resume_info fallback"
    # priority already used elsewhere (see phase3.py's linkedInUrl
    # resolution) — but only filling gaps, so a richer resume_info value
    # (e.g. a fuller skills list) is never clobbered by a thinner one.
    if consultant:
        def _blank(key: str) -> bool:
            val = resume_info.get(key)
            return val is None or (isinstance(val, str) and not val.strip())

        if _blank("linkedin") and consultant.linkedin_url:
            resume_info["linkedin"] = consultant.linkedin_url
        if _blank("phone") and consultant.phone:
            resume_info["phone"] = consultant.phone
        if _blank("full_name") and consultant.full_name:
            resume_info["full_name"] = consultant.full_name
        if _blank("email") and consultant.email:
            resume_info["email"] = consultant.email
        if _blank("title") and _blank("target_role") and _blank("target_title") and consultant.preferred_roles:
            resume_info["title"] = consultant.preferred_roles
        if (
            resume_info.get("years_experience") in (None, "")
            and resume_info.get("total_experience_years") in (None, "")
            and consultant.total_experience_years is not None
        ):
            resume_info["years_experience"] = float(consultant.total_experience_years)
        if not resume_info.get("skills") and not resume_info.get("tech_stack") and not resume_info.get("technical_proficiencies") and consultant.primary_skills:
            resume_info["skills"] = [s.strip() for s in consultant.primary_skills.split(",") if s.strip()]

    return resume_info


@router.get("/completeness")
async def get_resume_completeness(
    user_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """
    Pre-check the frontend calls before letting someone hit "Generate Resume".
    Returns which of skills/experience/education are still missing so the UI
    can show a checklist and disable the button, instead of the person
    finding out only after burning an AI call on a useless resume.
    """
    target_user_id, target_user = await _resolve_target_user(user_id, current_user, db)
    resume_info = await _build_resume_info(target_user, db)
    missing_fields = get_missing_resume_fields(resume_info)

    # Surface the target consultant's status so the "Generate/Regenerate
    # Resume" page (shared across admin/recruiter/consultant) can show a
    # clear "this consultant is inactive" indicator instead of the admin/
    # recruiter only finding out something's off after generating anyway.
    consultant_status = None
    consultant_res = await db.execute(select(Consultant.status).where(Consultant.user_id == target_user_id))
    consultant_status = consultant_res.scalar_one_or_none()

    return {
        "user_id": target_user_id,
        "complete": len(missing_fields) == 0,
        "missing_fields": missing_fields,
        "message": missing_fields_message(missing_fields) if missing_fields else None,
        "consultant_status": consultant_status,
        "is_active": bool(target_user.is_active),
    }


@router.post("/generate", response_model=ResumeResponse)
async def generate_resume(
    request: ResumeCreateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    target_user_id, target_user = await _resolve_target_user(request.user_id, current_user, db)
    resume_info = await _build_resume_info(target_user, db, request.experience_ids)

    # Guard rail: don't waste an AI call generating a resume from a profile
    # that's missing skills/experience/education — that just produces a
    # resume nobody can use. Fail fast with exactly what's missing so the
    # frontend (or the caller) can tell the consultant what to go fill in.
    missing_fields = get_missing_resume_fields(resume_info)
    if missing_fields:
        raise HTTPException(
            status_code=422,
            detail={
                "message": missing_fields_message(missing_fields),
                "missing_fields": missing_fields,
            },
        )

    try:
        generated_data, rate_limits, usage_info = generate_tailored_resume(resume_info, request.job_description or "General Role")
        if rate_limits:
            await save_claude_rate_limits(db, rate_limits)
        if usage_info:
            from phase8_ai_usage_service import log_ai_usage
            await log_ai_usage(
                db,
                purpose="resume_generation",
                model="claude-sonnet-4-6",
                input_tokens=usage_info["input_tokens"],
                output_tokens=usage_info["output_tokens"],
                consultant_id=str(target_user_id),
            )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Generation failed: {e}")
    # Compute a real ATS score from the generated resume vs. the job description.
    # Reuses phase6's keyword scorer and phase3's skill detector. Left as None
    # (no badge shown) when the JD has no recognizable skills to match against.
    ats_value = None
    try:
        from phase6 import _ats_score
        from phase3 import _detect_skills
        jd_skills = list(dict.fromkeys(
            _detect_skills(request.job_description or "")
            + (generated_data.get("missing_skills") or [])
        ))
        if jd_skills:
            text_parts = [
                generated_data.get("summary", ""),
                " ".join(generated_data.get("skills", []) or []),
                request.target_role or "",
            ]
            for exp in generated_data.get("experience", []) or []:
                text_parts.append(exp.get("role", ""))
                text_parts.append(" ".join(exp.get("bullets", []) or []))
            resume_text = " ".join(text_parts)
            ats_total, *_ = _ats_score(jd_skills, resume_text, request.target_role or "")
            ats_value = int(round(ats_total))
    except Exception as e:
        print(f"ATS scoring failed, leaving score empty: {e}")
        from error_logger import log_db_error
        await log_db_error(stage="ats_scoring", error=e)
        ats_value = None
    new_resume = Resume(
        user_id=target_user_id,
        requirement_id=request.requirement_id,
        title=request.title,
        target_role=request.target_role,
        job_description=request.job_description,
        data=generated_data,
        status='draft' if request.draft else 'generating',
        ats_score=ats_value,
    )

    db.add(new_resume)
    await db.commit()
    await db.refresh(new_resume)

    if request.draft:
        return new_resume

    # Generate DOCX and PDF using phase6 logic
    from phase6 import _generate_docx, _convert_to_pdf
    from pathlib import Path

    resume_dir = Path("/tmp/resumes") / str(target_user_id) / str(new_resume.id)
    resume_dir.mkdir(parents=True, exist_ok=True)

    docx_path = resume_dir / "resume.docx"
    pdf_path = resume_dir / "resume.pdf"

    try:
        _generate_docx(generated_data, docx_path)
        pdf_ok = _convert_to_pdf(docx_path, pdf_path)

        if pdf_ok:
            s3_key = f"users/{target_user_id}/resumes/{new_resume.id}/resume.pdf"
            with open(pdf_path, "rb") as f:
                if upload_file_to_s3(f, s3_key, "application/pdf"):
                    new_resume.s3_key = s3_key
                    new_resume.status = 'completed'
                else:
                    new_resume.status = 'failed_upload'
        else:
            new_resume.status = 'failed_pdf_conversion'

        # DOCX upload — was never wired up for this resume type before,
        # only the PDF ever got persisted. Key is derived from s3_key's
        # own naming pattern (same folder, extension swapped) so no new
        # column/migration is needed — the download endpoint derives the
        # same key back. Isolated in its own try/except on purpose: a
        # DOCX-only failure here must never downgrade the PDF status set
        # just above.
        try:
            docx_s3_key = f"users/{target_user_id}/resumes/{new_resume.id}/resume.docx"
            with open(docx_path, "rb") as f:
                upload_file_to_s3(
                    f, docx_s3_key,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as docx_exc:
            print(f"DOCX upload failed (PDF result unaffected): {docx_exc}")
            from error_logger import log_db_error
            await log_db_error(stage="resume_docx_s3_upload", error=docx_exc)
    except Exception as e:
        new_resume.status = 'failed_generation'
        print(f"Resume generation failed: {e}")
        from error_logger import log_db_error
        await log_db_error(
            stage="resume_generate_docx_pdf",
            error=e,
            source_type="resume",
            source_id=str(new_resume.id),
        )

    await db.commit()
    await db.refresh(new_resume)

    return new_resume

class FinalizeResumeRequest(BaseModel):
    data: dict

@router.post("/{resume_id}/finalize", response_model=ResumeResponse)
async def finalize_resume(
    resume_id: int,
    request: FinalizeResumeRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    resume = await _get_resume_for_user(db, resume_id, current_user)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    resume.data = request.data
    resume.status = 'generating'
    await db.commit()
    await db.refresh(resume)

    from phase6 import _generate_docx, _convert_to_pdf
    from pathlib import Path

    resume_dir = Path("/tmp/resumes") / str(resume.user_id) / str(resume.id)
    resume_dir.mkdir(parents=True, exist_ok=True)

    docx_path = resume_dir / "resume.docx"
    pdf_path = resume_dir / "resume.pdf"

    try:
        _generate_docx(resume.data, docx_path)
        pdf_ok = _convert_to_pdf(docx_path, pdf_path)

        if pdf_ok:
            s3_key = f"users/{resume.user_id}/resumes/{resume.id}/resume.pdf"
            with open(pdf_path, "rb") as f:
                if upload_file_to_s3(f, s3_key, "application/pdf"):
                    resume.s3_key = s3_key
                    resume.status = 'completed'
                else:
                    resume.status = 'failed_upload'
        else:
            resume.status = 'failed_pdf_conversion'

        # Same DOCX addition as generate_resume above — isolated so a
        # DOCX-only failure can't downgrade the PDF status set just above.
        try:
            docx_s3_key = f"users/{resume.user_id}/resumes/{resume.id}/resume.docx"
            with open(docx_path, "rb") as f:
                upload_file_to_s3(
                    f, docx_s3_key,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as docx_exc:
            print(f"DOCX upload failed (PDF result unaffected): {docx_exc}")
            from error_logger import log_db_error
            await log_db_error(stage="resume_finalize_docx_s3_upload", error=docx_exc)
    except Exception as e:
        resume.status = 'failed_generation'
        print(f"Resume generation failed: {e}")
        from error_logger import log_db_error
        await log_db_error(
            stage="resume_finalize_docx_pdf",
            error=e,
            source_type="resume",
            source_id=str(resume.id),
        )

    await db.commit()
    await db.refresh(resume)

    return resume

@router.post("/upload", response_model=ResumeResponse)
async def upload_resume(
    title: str = Form(...),
    target_role: str = Form(None),
    # BUG FIX: this endpoint always saved the resume under the UPLOADER's
    # own account (current_user.id) — so when an admin/recruiter selected
    # a candidate on the My Resumes page and uploaded a PDF "for" them, it
    # actually got attached to the admin/recruiter's own account instead.
    # The upload itself succeeded (hence the success toast), but the
    # candidate's resume list is filtered by user_id, so it could never
    # show up there — looked exactly like "upload succeeds, nothing
    # appears". Now accepts an optional target_user_id so the frontend can
    # say who it's actually for.
    target_user_id: int = Form(None),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    # CHANGE: was PDF-only ("Only PDF files are supported."). Now DOCX-only
    # instead — content-type and storage key updated to match.
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")

    owner_id = current_user.id
    if target_user_id and target_user_id != current_user.id:
        if current_user.role not in ("ADMIN", "RECRUITER"):
            raise HTTPException(status_code=403, detail="Only admins/recruiters can upload a resume for someone else.")
        if current_user.role == "RECRUITER":
            allowed = await db.execute(
                select(Consultant.user_id).where(
                    Consultant.user_id == target_user_id,
                    or_(
                        Consultant.sales_recruiter_user_id == current_user.id,
                        Consultant.id.in_(
                            select(RecruiterConsultant.consultant_id).where(
                                RecruiterConsultant.recruiter_id == current_user.id
                            )
                        ),
                    ),
                )
            )
            if not allowed.scalars().first():
                raise HTTPException(status_code=403, detail="That consultant isn't assigned to you.")
        owner_id = target_user_id

    s3_key = f"users/{owner_id}/resumes/{uuid.uuid4()}/final.docx"

    # Read into memory once so the same bytes can both upload to Spaces AND
    # feed text extraction below — file.file's stream position isn't safe
    # to rely on after upload_fileobj() consumes it.
    file_bytes = await file.read()

    success = upload_file_to_s3(
        io.BytesIO(file_bytes),
        s3_key,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    if not success:
        raise HTTPException(status_code=500, detail="Failed to upload file to S3.")

    # BUG FIX: manually-uploaded resumes previously had no extracted
    # content at all — `data` stayed the JSONB column's empty-dict default
    # forever. The "Edit" button in My Resumes always routed every non-base
    # resume to the structured JSON editor (personal info / experience /
    # education fields), which had nothing to populate those fields with,
    # so it silently rendered a blank template with just the label showing
    # — looked broken with no error anywhere. Extracting text here (same
    # DOCX -> text logic the base resume already uses) and storing it under
    # data["raw_text"] gives these resumes something real to edit, via a
    # plain-text editor instead of the structured one — see GET/PUT
    # /{id}/text below.
    extracted_text = _extract_text_from_docx(file_bytes)

    new_resume = Resume(
        user_id=owner_id,
        title=title,
        target_role=target_role,
        s3_key=s3_key,
        status='completed',
        data={"raw_text": extracted_text},
    )

    db.add(new_resume)
    await db.commit()
    await db.refresh(new_resume)

    return new_resume

@router.get("/list", response_model=PaginatedResumes)
async def list_resumes(
    page: int = 1,
    page_size: int = 10,
    search: Optional[str] = None,
    user_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    if current_user.role == "ADMIN":
        query = select(Resume)
    elif current_user.role == "RECRUITER":
        consultant_users_query = select(Consultant.user_id).where(
            or_(
                Consultant.sales_recruiter_user_id == current_user.id,
                Consultant.id.in_(
                    select(RecruiterConsultant.consultant_id).where(
                        RecruiterConsultant.recruiter_id == current_user.id
                    )
                )
            )
        )
        query = select(Resume).where(
            or_(
                Resume.user_id == current_user.id,
                Resume.user_id.in_(consultant_users_query)
            )
        )
    else:
        query = select(Resume).where(Resume.user_id == current_user.id)

    if user_id:
        # Additional safety to only allow filtering if they have access
        if current_user.role == "ADMIN" or current_user.role == "RECRUITER":
            query = query.where(Resume.user_id == user_id)

    if search:
        query = query.where(Resume.title.ilike(f"%{search}%"))

    count_query = select(func.count()).select_from(query.subquery())
    total = (await db.execute(count_query)).scalar_one()

    query = query.order_by(Resume.created_at.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)

    resumes = (await db.execute(query)).scalars().all()
    total_pages = math.ceil(total / page_size) if page_size > 0 else 0

    # generate s3_url for each resume
    response_data = []
    for r in resumes:
        r_dict = {
            "id": r.id,
            "user_id": r.user_id,
            "requirement_id": r.requirement_id,
            "title": r.title,
            "target_role": r.target_role,
            "job_description": r.job_description,
            "data": r.data or {},
            "s3_key": r.s3_key,
            "s3_url": generate_presigned_url(r.s3_key) if r.s3_key else None,
            "ats_score": r.ats_score,
            "status": r.status,
            "download_count": r.download_count,
            "last_downloaded": r.last_downloaded,
            "created_at": r.created_at,
            "updated_at": r.updated_at
        }
        response_data.append(ResumeResponse(**r_dict))

    # --- Pin the consultant's base resume (from their profile) at the top ---
    # It lives on the consultants table + Spaces, not the resumes table, so we surface it
    # here as a read-only entry instead of duplicating a row (single source of truth).
    if page == 1:
        base_target_user_id = None
        if current_user.role == "CONSULTANT":
            base_target_user_id = current_user.id
        elif user_id:  # admin/recruiter viewing a specific candidate
            base_target_user_id = user_id

        if base_target_user_id is not None and (not search or search.lower() in "base resume"):
            base_consultant = (await db.execute(
                select(Consultant).where(Consultant.user_id == base_target_user_id)
            )).scalar_one_or_none()

            # FEATURE CHANGE: base resume is now generated from the
            # profile (see sync_base_resume_text in this file, called by
            # phase3.py on every profile/experience save) rather than
            # uploaded — base_resume_file_path is a legacy field from the
            # old upload flow and stays NULL forever for every consultant
            # going forward, so gating on it alone hid this card for
            # everyone. Show it whenever the consultant profile exists at
            # all; an empty/just-started profile still has a valid (if
            # sparse) generated resume to view/edit.
            if base_consultant:
                base_ts = base_consultant.updated_at or datetime.now(timezone.utc)
                response_data.insert(0, ResumeResponse(
                    id=-1,
                    user_id=base_target_user_id,
                    requirement_id=None,
                    title="Base Resume",
                    target_role="From profile",
                    job_description=None,
                    data={},
                    s3_key=None,
                    s3_url=None,  # resolved on demand via GET /api/resume/base/download
                    ats_score=None,
                    status="base",
                    download_count=0,
                    last_downloaded=None,
                    created_at=base_ts,
                    updated_at=base_ts,
                    is_base=True,
                ))
                total += 1

    return PaginatedResumes(
        data=response_data,
        total=total,
        page=page,
        page_size=page_size,
        total_pages=total_pages
    )

class BaseResumeTextDTO(BaseModel):
    text: Optional[str] = ""
    filename: Optional[str] = None

class BaseResumeTextUpdateRequest(BaseModel):
    text: str

@router.get("/base/text", response_model=BaseResumeTextDTO)
async def get_base_resume_text(
    user_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    if current_user.role == "CONSULTANT":
        target_user_id = current_user.id
    else:
        target_user_id = user_id or current_user.id
        
    resolved_id, target_user = await _resolve_target_user(target_user_id, current_user, db)
    
    consultant = (await db.execute(
        select(Consultant).where(Consultant.user_id == resolved_id)
    )).scalar_one_or_none()
    
    if not consultant:
        raise HTTPException(status_code=404, detail="Consultant profile not found")
        
    from pathlib import Path
    filename = Path(consultant.base_resume_file_path).name if consultant.base_resume_file_path else None
    
    return BaseResumeTextDTO(
        text=consultant.base_resume_text or "",
        filename=filename
    )

@router.put("/base/text")
async def update_base_resume_text(
    request: BaseResumeTextUpdateRequest,
    user_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    if current_user.role == "CONSULTANT":
        target_user_id = current_user.id
    else:
        target_user_id = user_id or current_user.id
        
    resolved_id, target_user = await _resolve_target_user(target_user_id, current_user, db)
    
    consultant = (await db.execute(
        select(Consultant).where(Consultant.user_id == resolved_id)
    )).scalar_one_or_none()
    
    if not consultant:
        raise HTTPException(status_code=404, detail="Consultant profile not found")
        
    consultant.base_resume_text = request.text
    # The structured JSON editor (GET/PUT /base/content) must not keep
    # showing whatever was parsed before this plain-text edit — clear it
    # so its backfill re-parses the NEW text next time that page loads
    # (same reasoning as the upload endpoint in phase3.py).
    consultant.base_resume_content = None
    # BUG FIX: this endpoint only ever updated base_resume_text (the plain
    # text used for AI tailoring/matching) — the actual file streamed back
    # by GET /base/download reads consultant.base_resume_file_path, which
    # was never touched here. Edits saved on this page had no effect on
    # what "Download" served; it kept returning the exact original
    # uploaded file forever. Base resumes are DOCX-only (see phase3.py's
    # ALLOWED_RESUME_TYPES), so regenerate a DOCX from the edited text and
    # overwrite the file in place — same path/filename, so every other
    # place that already stores or serves base_resume_file_path keeps
    # working unchanged.
    import io
    from pathlib import Path
    from docx import Document
    try:
        doc = Document()
        for line in (request.text or "").split("\n"):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        stored = consultant.base_resume_file_path
        if stored and os.path.isfile(stored):
            # Locally stored file — how every current upload is stored (see
            # phase3.py's _save_resume_file). Some records predate the
            # DOCX-only upload restriction and still end in .pdf — force
            # the extension to .docx and drop the old file rather than
            # writing DOCX bytes under a stale .pdf name, which would
            # mislabel the media_type the download endpoint serves.
            old_path = Path(stored)
            new_path = old_path.with_suffix(".docx")
            new_path.parent.mkdir(parents=True, exist_ok=True)
            new_path.write_bytes(docx_bytes)
            if new_path != old_path and old_path.exists():
                old_path.unlink()
            consultant.base_resume_file_path = str(new_path)
        elif stored:
            # Not a local path -> a Spaces/S3 object key (mirrors the same
            # local-vs-Spaces branching download_base_resume already does).
            # Re-upload under a .docx key so a pre-migration .pdf key
            # doesn't end up mislabeled the same way as the local case.
            key = stored if stored.lower().endswith(".docx") else str(Path(stored).with_suffix(".docx"))
            if upload_file_to_s3(
                io.BytesIO(docx_bytes), key,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ):
                consultant.base_resume_file_path = key
        else:
            # No file on record yet (text entered without ever uploading a
            # file) — create one locally, same layout _save_resume_file uses.
            upload_dir = Path(os.getenv("UPLOAD_DIR", "uploads/resumes")) / str(consultant.id)
            upload_dir.mkdir(parents=True, exist_ok=True)
            new_path = upload_dir / f"{uuid.uuid4().hex}.docx"
            new_path.write_bytes(docx_bytes)
            consultant.base_resume_file_path = str(new_path)
    except Exception as e:
        # Don't fail the save over a DOCX regen hiccup — base_resume_text
        # itself already committed below and still powers AI tailoring.
        print(f"Base resume DOCX regeneration on save failed for consultant {consultant.id}: {e}")
        from error_logger import log_db_error
        await log_db_error(
            stage="base_resume_docx_regen_on_save",
            error=e,
            source_type="consultant",
            source_id=str(consultant.id),
        )
    await db.commit()
    return {"success": True, "message": "Base resume text updated successfully"}


class BaseResumeContentDTO(BaseModel):
    content: dict
    filename: Optional[str] = None


class BaseResumeContentUpdateRequest(BaseModel):
    content: dict


_HEURISTIC_SECTION_HEADERS = [
    ("career_objective", re.compile(r'^career\s+objective\s*:?$', re.I)),
    ("technical_proficiencies", re.compile(r'^technical\s+proficienc(?:y|ies)\s*:?$', re.I)),
    ("experience", re.compile(r'^(?:experience|work\s+experience|professional\s+experience)\s*:?$', re.I)),
    ("educational_background", re.compile(r'^educational\s+background\s*:?$|^education\s*:?$', re.I)),
]
_HEURISTIC_ASSOCIATED_RE = re.compile(r'^Associated with\s+(.+?)\s*\(([^)]*)\)\s*$', re.I)
_HEURISTIC_DESIGNATION_RE = re.compile(r'^Designation:\s*(.+)$', re.I)
_HEURISTIC_EDU_LINE_RE = re.compile(
    r'(university|college|institute|bachelor|master|b\.?tech|m\.?tech|b\.?s\.?c?\b|m\.?s\.?c?\b|ph\.?d|diploma).*\(\d{4}\)\s*$',
    re.I,
)


def _heuristic_parse_resume_text(text: str) -> Optional[dict]:
    """
    Free, non-AI fallback for turning base_resume_text into the same
    structured shape parse_resume_text_to_structured_data produces —
    used when the Claude call is unavailable (e.g. billing/credits, no
    API key). Only understands the specific template these base resumes
    follow: Name / contact line / "CAREER OBJECTIVE:" / "TECHNICAL
    PROFICIENCIES:" (a table, so its actual content — regardless of
    where the header sits — is appended after every other paragraph by
    phase3.py's _extract_text_from_docx; see the pairs picked up at the
    end below) / "EXPERIENCE:" ("Associated with X (dates)" +
    "Designation: Y" + bullet lines per role) / "EDUCATIONAL
    BACKGROUND:". Returns None if the text doesn't contain any of these
    headers, so the caller knows to leave base_resume_content unset
    (and retry via AI later) rather than saving a bad guess.
    """
    if not text or not text.strip():
        return None
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if len(lines) < 3:
        return None

    header_positions = []
    for idx, line in enumerate(lines):
        for key, pattern in _HEURISTIC_SECTION_HEADERS:
            if pattern.match(line):
                header_positions.append((key, idx))
                break
    if not header_positions:
        return None

    first_header_idx = header_positions[0][1]
    name = lines[0] if first_header_idx > 0 else "Unknown"
    contact_line = lines[1] if first_header_idx > 1 else ""

    email = ""
    phone = ""
    linkedin = ""
    if contact_line:
        m = re.search(r'[\w.\-+]+@[\w.\-]+\.\w+', contact_line)
        if m:
            email = m.group(0)
        m = re.search(r'(\+?\d[\d\s\-().]{7,}\d)', contact_line)
        if m:
            phone = m.group(1).strip()
        m = re.search(r'https?://\S+', contact_line)
        if m:
            linkedin = m.group(0).rstrip('.,')

    boundaries = header_positions + [("__end__", len(lines))]
    sections: dict = {}
    for i, (key, start) in enumerate(header_positions):
        _, end = boundaries[i + 1]
        sections[key] = lines[start + 1:end]

    career_objective = " ".join(sections.get("career_objective", [])).strip()

    experience = []
    current = None
    for line in sections.get("experience", []):
        m = _HEURISTIC_ASSOCIATED_RE.match(line)
        if m:
            if current:
                experience.append(current)
            dates = m.group(2).strip()
            start_date, sep, end_date = dates.partition("–")
            if not sep:
                start_date, sep, end_date = dates.partition("-")
            current = {
                "client": m.group(1).strip(),
                "role": "",
                # ResumeRichPreview/ResumeEditorPage's ResumeData type reads
                # exp.start/exp.end (not start_date/end_date) — using the
                # wrong keys here meant dates parsed fine but silently never
                # rendered anywhere.
                "start": start_date.strip(),
                "end": end_date.strip() if sep else "",
                "location": "",
                "bullets": [],
            }
            continue
        m = _HEURISTIC_DESIGNATION_RE.match(line)
        if m and current is not None:
            current["role"] = m.group(1).strip()
            continue
        if current is not None:
            current["bullets"].append(line)
    if current:
        experience.append(current)

    # Whatever's left after the last recognized section's content is the
    # flattened technical-proficiencies table (see docstring above). When
    # "Educational Background" is that last section there's no header to
    # bound where the degree entries end and the table begins, so only
    # lines that actually look like a degree entry (mentions a
    # university/degree and ends in a "(YYYY)") are kept as education;
    # everything after the run of those is treated as trailing table
    # content instead of being swallowed into education.
    last_key, last_start = header_positions[-1]
    last_section_lines = sections.get(last_key, [])

    if last_key == "educational_background":
        education = []
        consumed = 0
        for line in last_section_lines:
            if _HEURISTIC_EDU_LINE_RE.search(line):
                education.append({"degree": line, "institution": "", "year": ""})
                consumed += 1
            else:
                break
        if not education and last_section_lines:
            education.append({"degree": last_section_lines[0], "institution": "", "year": ""})
            consumed = 1
        trailing = last_section_lines[consumed:]
    else:
        education = [
            {"degree": line, "institution": "", "year": ""}
            for line in sections.get("educational_background", [])
        ]
        trailing = last_section_lines

    technical_proficiencies = []
    if trailing and len(trailing) % 2 == 0:
        for i in range(0, len(trailing), 2):
            technical_proficiencies.append({"category": trailing[i], "skills": trailing[i + 1]})
    elif trailing:
        technical_proficiencies.append({"category": "Skills", "skills": ", ".join(trailing)})

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "location": "",
        "linkedin": linkedin,
        "github": "",
        "career_objective": career_objective,
        "summary": "",
        "technical_proficiencies": technical_proficiencies,
        "skills": [],
        "experience": experience,
        "education": education,
        "certifications": [],
        "generation_notes": "Parsed automatically from plain text (no AI used) — please review for accuracy.",
    }


def _flatten_base_resume_content_to_text(data: dict) -> str:
    """
    Plain-text projection of the structured base resume content, kept in
    base_resume_text so nothing that already reads it (AI tailoring in
    claude_service.py, candidate matching in matching_router.py) has to
    change or even know this structured editor exists.
    """
    if not data:
        return ""
    parts: List[str] = []

    def add(label: Optional[str], value):
        if value is None:
            return
        if isinstance(value, (list, tuple)):
            value = ", ".join(str(v) for v in value if v)
        value = str(value).strip()
        if not value:
            return
        parts.append(f"{label}: {value}" if label else value)

    add(None, data.get("name"))
    add("Career Objective", data.get("career_objective") or data.get("summary"))

    tech = data.get("technical_proficiencies")
    if isinstance(tech, list) and tech:
        for tp in tech:
            skills = tp.get("skills") if isinstance(tp, dict) else None
            add(tp.get("category") if isinstance(tp, dict) else None, skills)
    else:
        add("Skills", data.get("skills"))

    for exp in data.get("experience") or []:
        if not isinstance(exp, dict):
            continue
        header = " - ".join(
            str(v) for v in [exp.get("client") or exp.get("company"), exp.get("role") or exp.get("title")] if v
        )
        add(None, header)
        add(None, exp.get("description"))
        add(None, exp.get("bullets"))

    for proj in (data.get("key_projects") or []) + (data.get("academic_projects") or []):
        if not isinstance(proj, dict):
            continue
        add(None, proj.get("title") or proj.get("game_name"))
        add(None, proj.get("description"))
        add(None, proj.get("responsibilities"))

    edu = data.get("education") or data.get("educational_background")
    if isinstance(edu, list):
        for e in edu:
            if isinstance(e, dict):
                add(None, " - ".join(str(v) for v in [e.get("degree"), e.get("institution") or e.get("college"), e.get("year")] if v))
            else:
                add(None, e)
    elif edu:
        add(None, edu)

    add("Certifications", data.get("certifications"))
    add("Non-Technical Proficiencies", data.get("non_technical_proficiencies"))
    add("Achievements", data.get("achievements"))

    return "\n".join(parts)


_MONTH_YEAR_FORMATS = ("%b %Y", "%B %Y")


def _parse_month_year(value: Optional[str]) -> Optional[date]:
    """Parse the Base Resume editor's free-text 'Mon YYYY' / 'Month YYYY'
    fields into a real date (day fixed to 1). Returns None on anything
    unparseable rather than raising."""
    if not value or not value.strip():
        return None
    cleaned = value.strip()
    for fmt in _MONTH_YEAR_FORMATS:
        try:
            return datetime.strptime(cleaned, fmt).date().replace(day=1)
        except ValueError:
            continue
    return None


def _parse_resume_editor_dates(start: Optional[str], end: Optional[str]):
    start_date = _parse_month_year(start)
    end_stripped = (end or "").strip().lower()
    if not end_stripped or end_stripped in ("present", "current"):
        return start_date, None, True
    return start_date, _parse_month_year(end), False


def _format_month_year(d: Optional[date]) -> str:
    return d.strftime("%b %Y") if d else ""


async def build_base_resume_content(db: AsyncSession, consultant: Consultant) -> dict:
    """Overlay the profile-derived fields (Skills, Experience, Education,
    Name/Phone/Location/LinkedIn) fresh from Consultant/ConsultantExperience/
    User.resume_info onto whatever resume-only extras (Certifications,
    career objective wording, etc.) were previously saved. This is the
    single source of truth both GET /base/content and the auto-sync below
    use — nothing here is ever read from a stale cached blob."""
    extra_content = dict(consultant.base_resume_content or {})
    for _stale_key in (
        "name", "email", "phone", "location", "linkedin",
        "career_objective", "summary", "technical_proficiencies",
        "experience", "education",
    ):
        extra_content.pop(_stale_key, None)

    resume_info: dict = {}
    if consultant.user_id:
        info_result = await db.execute(select(User.resume_info).where(User.id == consultant.user_id))
        resume_info = info_result.scalar_one_or_none() or {}
    summary_text = resume_info.get("summary") or ""

    technical_proficiencies = []
    if (consultant.primary_skills or "").strip():
        technical_proficiencies.append({"category": "Primary Skills", "skills": consultant.primary_skills})
    if (consultant.secondary_skills or "").strip():
        technical_proficiencies.append({"category": "Secondary Skills", "skills": consultant.secondary_skills})

    exp_rows_result = await db.execute(
        select(ConsultantExperience)
        .where(ConsultantExperience.consultant_id == consultant.id)
        .order_by(ConsultantExperience.sort_order.asc())
    )
    experience_list = []
    for exp in exp_rows_result.scalars().all():
        bullets = [b for b in (exp.responsibilities, exp.achievements) if b]
        experience_list.append({
            "id": str(exp.id),
            "role": exp.role_title or "",
            "client": exp.client_name or "",
            "start": _format_month_year(exp.start_date),
            "end": "Present" if exp.is_present else _format_month_year(exp.end_date),
            "location": exp.location or "",
            "bullets": bullets,
            "technologies": exp.technologies or [],
        })

    return {
        **extra_content,
        "name": consultant.full_name or "",
        "email": consultant.email or "",
        "phone": consultant.phone or "",
        "location": consultant.current_location or "",
        "linkedin": consultant.linkedin_url or "",
        "career_objective": summary_text,
        "summary": summary_text,
        "technical_proficiencies": technical_proficiencies,
        "experience": experience_list,
        "education": consultant.education or [],
    }


async def sync_base_resume_text(db: AsyncSession, consultant: Consultant) -> None:
    """Regenerate and stage consultant.base_resume_text from CURRENT
    profile fields — called by phase3.py whenever Skills, Experience,
    Education, or contact info changes on My Profile, so base_resume_text
    (read by both the completeness check and matching_router.py's TF-IDF
    scoring) never sits stale waiting for someone to manually open and
    save the Base Resume editor. Stages the change on consultant only —
    caller still owns db.commit()."""
    content = await build_base_resume_content(db, consultant)
    consultant.base_resume_text = _flatten_base_resume_content_to_text(content)


@router.get("/base/content", response_model=BaseResumeContentDTO)
async def get_base_resume_content(
    user_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Structured counterpart to GET /base/text — powers the rich JSON + preview
    Base Resume editor (mirrors GET .../resume/content in phase6.py for
    tailored resumes). Same auth/resolution as /base/text.
    """
    if current_user.role == "CONSULTANT":
        target_user_id = current_user.id
    else:
        target_user_id = user_id or current_user.id

    resolved_id, target_user = await _resolve_target_user(target_user_id, current_user, db)

    consultant = (await db.execute(
        select(Consultant).where(Consultant.user_id == resolved_id)
    )).scalar_one_or_none()

    if not consultant:
        raise HTTPException(status_code=404, detail="Consultant profile not found")

    from pathlib import Path
    filename = Path(consultant.base_resume_file_path).name if consultant.base_resume_file_path else None

    # BACKFILL: consultants whose base resume predates this structured
    # editor only have base_resume_text (plain text) — base_resume_content
    # is null, which would show a blank {} editor even though a real
    # resume exists.
    #
    # ORDER FIX: this used to try the AI parser first and only fall back
    # to the free heuristic parser after it failed — meaning every page
    # load for a not-yet-backfilled consultant blocked on a real network
    # call to Anthropic (up to the 30s timeout below) before the page
    # could render at all. That's exactly why loading felt inconsistent
    # ("sometimes opens, sometimes stuck") — fast for already-backfilled
    # consultants, but genuinely slow (not actually hung) for ones still
    # waiting on that AI round-trip. Since these base resumes reliably
    # follow one template, try the free, synchronous, near-instant
    # heuristic parser FIRST — the common case now resolves immediately
    # — and only reach for the AI parser as a fallback when the text
    # doesn't match that template at all.
    if not consultant.base_resume_content and (consultant.base_resume_text or "").strip():
        try:
            heuristic_data = _heuristic_parse_resume_text(consultant.base_resume_text)
            if heuristic_data:
                consultant.base_resume_content = heuristic_data
                await db.commit()
        except Exception as e:
            print(f"Base resume content backfill (heuristic) failed for consultant {consultant.id}: {e}")

    if not consultant.base_resume_content and (consultant.base_resume_text or "").strip():
        from claude_service import parse_resume_text_to_structured_data
        try:
            # asyncio.to_thread runs the blocking Anthropic call on a
            # worker thread so it doesn't freeze the event loop (and
            # every other request the backend is serving) while it's in
            # flight — see claude_service.py for the 30s timeout that
            # bounds how long this can take.
            parsed_data, rate_limits, usage_info = await asyncio.to_thread(
                parse_resume_text_to_structured_data, consultant.base_resume_text
            )
            if rate_limits:
                await save_claude_rate_limits(db, rate_limits)
            if usage_info:
                # usage_info is only non-None when the AI call actually
                # succeeded — only THEN persist the result. Otherwise
                # parsed_data is just its blank "Automatic parsing was
                # unavailable..." skeleton (e.g. ANTHROPIC_API_KEY
                # missing/invalid, no credits, or a transient API error)
                # — saving that would look identical to a completed
                # backfill and permanently block every future retry.
                # Leave it unsaved so the next GET tries again.
                from phase8_ai_usage_service import log_ai_usage
                await log_ai_usage(
                    db,
                    purpose="base_resume_backfill",
                    model="claude-sonnet-4-6",
                    input_tokens=usage_info["input_tokens"],
                    output_tokens=usage_info["output_tokens"],
                    consultant_id=str(resolved_id),
                )
                consultant.base_resume_content = parsed_data
                await db.commit()
            else:
                print(f"Base resume content backfill returned no usage_info (AI call unavailable) for consultant {consultant.id}.")
        except Exception as e:
            # Don't fail the read over a backfill hiccup — the next GET
            # will just retry both parsers.
            print(f"Base resume content backfill (AI) failed for consultant {consultant.id}: {e}")
            from error_logger import log_db_error
            await log_db_error(
                stage="base_resume_content_backfill",
                error=e,
                source_type="consultant",
                source_id=str(consultant.id),
            )

    content = await build_base_resume_content(db, consultant)
    return BaseResumeContentDTO(
        content=content,
        filename=filename,
    )


@router.put("/base/content")
async def update_base_resume_content(
    request: BaseResumeContentUpdateRequest,
    user_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Structured counterpart to PUT /base/text (same bug-fix shape: the DOCX
    served by GET /base/download must stay in sync with what was last
    saved here). Regenerates the DOCX via phase6.py's _generate_docx —
    the same builder tailored/uploaded resumes use — so the downloaded
    file matches the ResumeRichPreview format the user edits against.
    Also re-derives base_resume_text so AI tailoring/matching keep working
    unchanged; neither base_resume_text nor base_resume_file_path's
    resolution logic elsewhere is touched.
    """
    if current_user.role == "CONSULTANT":
        target_user_id = current_user.id
    else:
        target_user_id = user_id or current_user.id

    resolved_id, target_user = await _resolve_target_user(target_user_id, current_user, db)

    consultant = (await db.execute(
        select(Consultant).where(Consultant.user_id == resolved_id)
    )).scalar_one_or_none()

    if not consultant:
        raise HTTPException(status_code=404, detail="Consultant profile not found")

    resume_data = dict(request.content or {})

    if (resume_data.get("name") or "").strip():
        consultant.full_name = resume_data["name"].strip()
    if resume_data.get("phone") is not None:
        consultant.phone = resume_data.get("phone") or None
    if resume_data.get("location") is not None:
        consultant.current_location = resume_data.get("location") or None
    if resume_data.get("linkedin") is not None:
        consultant.linkedin_url = resume_data.get("linkedin") or None
    if "education" in resume_data:
        consultant.education = resume_data.get("education") or []

    summary_val = resume_data.get("career_objective")
    if summary_val is None:
        summary_val = resume_data.get("summary")
    if summary_val is not None and consultant.user_id:
        user_row_result = await db.execute(select(User).where(User.id == consultant.user_id))
        user_row = user_row_result.scalar_one_or_none()
        if user_row:
            info = dict(user_row.resume_info or {})
            info["summary"] = summary_val
            user_row.resume_info = info

    tech_rows = resume_data.get("technical_proficiencies") or []
    primary_bits, secondary_bits = [], []
    for row in tech_rows:
        if not isinstance(row, dict):
            continue
        skills_val = row.get("skills")
        skills_str = ", ".join(skills_val) if isinstance(skills_val, list) else (skills_val or "")
        skills_str = skills_str.strip()
        if not skills_str:
            continue
        category = (row.get("category") or "").strip().lower()
        if "primary" in category:
            primary_bits.append(skills_str)
        elif "secondary" in category:
            secondary_bits.append(skills_str)
        else:
            secondary_bits.append(skills_str)
    if primary_bits:
        consultant.primary_skills = ", ".join(primary_bits)
    if secondary_bits:
        consultant.secondary_skills = ", ".join(secondary_bits)

    incoming_experience = resume_data.get("experience") or []
    existing_rows_result = await db.execute(
        select(ConsultantExperience).where(ConsultantExperience.consultant_id == consultant.id)
    )
    existing_by_id = {e.id: e for e in existing_rows_result.scalars().all()}
    seen_ids = set()
    reconciled_experience = []

    for idx, item in enumerate(incoming_experience):
        if not isinstance(item, dict):
            continue
        raw_id = item.get("id")
        exp_id = None
        if raw_id not in (None, "", "new"):
            try:
                exp_id = int(raw_id)
            except (TypeError, ValueError):
                exp_id = None

        start_date, end_date, is_present = _parse_resume_editor_dates(item.get("start"), item.get("end"))
        bullets = item.get("bullets") or []
        responsibilities = bullets[0] if len(bullets) > 0 else None
        achievements = "\n".join(bullets[1:]) if len(bullets) > 1 else None
        technologies = item.get("technologies") or []

        if exp_id is not None and exp_id in existing_by_id:
            exp = existing_by_id[exp_id]
            exp.client_name = item.get("client") or exp.client_name
            exp.role_title = item.get("role") or exp.role_title
            exp.location = item.get("location")
            if start_date is not None:
                exp.start_date = start_date
            exp.end_date = end_date
            exp.is_present = is_present
            exp.technologies = technologies
            exp.responsibilities = responsibilities
            exp.achievements = achievements
            exp.sort_order = idx
            seen_ids.add(exp_id)
            reconciled_experience.append(item)
        else:
            new_exp = ConsultantExperience(
                consultant_id=consultant.id,
                client_name=item.get("client") or "",
                role_title=item.get("role") or "",
                start_date=start_date or date.today().replace(day=1),
                end_date=end_date,
                is_present=is_present,
                location=item.get("location"),
                technologies=technologies,
                responsibilities=responsibilities,
                achievements=achievements,
                sort_order=idx,
            )
            db.add(new_exp)
            await db.flush()
            seen_ids.add(new_exp.id)
            reconciled_experience.append({**item, "id": str(new_exp.id)})

    for old_id, old_exp in existing_by_id.items():
        if old_id not in seen_ids:
            await db.delete(old_exp)

    resume_data["experience"] = reconciled_experience

    consultant.base_resume_content = resume_data
    consultant.base_resume_text = _flatten_base_resume_content_to_text(resume_data)

    import io
    import tempfile
    from pathlib import Path
    from phase6 import _generate_docx
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            docx_path = Path(tmp_dir) / "base_resume.docx"
            _generate_docx(resume_data, docx_path)
            docx_bytes = docx_path.read_bytes()

        stored = consultant.base_resume_file_path
        if stored and os.path.isfile(stored):
            old_path = Path(stored)
            new_path = old_path.with_suffix(".docx")
            new_path.parent.mkdir(parents=True, exist_ok=True)
            new_path.write_bytes(docx_bytes)
            if new_path != old_path and old_path.exists():
                old_path.unlink()
            consultant.base_resume_file_path = str(new_path)
        elif stored:
            key = stored if stored.lower().endswith(".docx") else str(Path(stored).with_suffix(".docx"))
            if upload_file_to_s3(
                io.BytesIO(docx_bytes), key,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ):
                consultant.base_resume_file_path = key
        else:
            upload_dir = Path(os.getenv("UPLOAD_DIR", "uploads/resumes")) / str(consultant.id)
            upload_dir.mkdir(parents=True, exist_ok=True)
            new_path = upload_dir / f"{uuid.uuid4().hex}.docx"
            new_path.write_bytes(docx_bytes)
            consultant.base_resume_file_path = str(new_path)
    except Exception as e:
        print(f"Base resume DOCX regeneration on content save failed for consultant {consultant.id}: {e}")
        from error_logger import log_db_error
        await log_db_error(
            stage="base_resume_docx_regen_on_content_save",
            error=e,
            source_type="consultant",
            source_id=str(consultant.id),
        )

    # BUG FIX: this was indented inside the `except` above, so it only
    # ran when DOCX regeneration failed — on the normal success path
    # every profile/experience field written earlier in this function
    # (name, phone, location, LinkedIn, education, skills, the full
    # experience reconciliation, base_resume_content, base_resume_text)
    # was silently discarded, never reaching the database, even though
    # the response below still claimed success. Commit needs to happen
    # unconditionally — whether or not the DOCX step itself succeeded.
    await db.commit()
    return {"success": True, "message": "Base resume content updated successfully"}


@router.get("/{id}/text", response_model=BaseResumeTextDTO)
async def get_resume_text(
    id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Plain-text counterpart to the structured GET /{id} — for a manually
    uploaded resume (see /upload above), `data` only ever holds
    {"raw_text": "..."}, never the structured personal_info/experience/
    education fields the JSON editor (GET /{id} + ResumeEditorPage.tsx)
    expects. Reusing that editor for these resumes rendered a blank
    template with every field empty and no error anywhere — this endpoint
    (and PUT below) back a dedicated plain-text editor instead, mirroring
    GET/PUT /base/text for the base resume.
    """
    resume = await _get_resume_for_user(db, id, current_user)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    return BaseResumeTextDTO(
        text=(resume.data or {}).get("raw_text", "") or "",
        filename=resume.title,
    )


@router.put("/{id}/text")
async def update_resume_text(
    id: int,
    request: BaseResumeTextUpdateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    resume = await _get_resume_for_user(db, id, current_user)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Reassign (rather than mutate in place) so SQLAlchemy's change
    # tracking reliably picks up the JSONB column update — mutating a
    # dict in place on some backends/drivers doesn't always mark the
    # attribute dirty for JSON columns.
    existing = dict(resume.data or {})
    existing["raw_text"] = request.text
    resume.data = existing
    # BUG FIX: this only ever updated data["raw_text"] — the actual file
    # streamed back by View (/{id}/view) and Download (/{id}/download,
    # /{id}/download/file) reads resume.s3_key, which was never touched
    # here. Saving an edit had no effect on what View/Download served; it
    # kept returning the exact original uploaded file forever — same class
    # of bug as update_base_resume_text above, same fix: regenerate a DOCX
    # from the edited text and overwrite the object at the same s3_key, so
    # nothing else that already references s3_key needs to change.
    if resume.s3_key:
        import io
        from docx import Document
        try:
            doc = Document()
            for line in (request.text or "").split("\n"):
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()
            # Uploads always write s3_key ending in .docx (see /upload
            # above) — force it here too in case this endpoint is ever
            # reached for an older/differently-keyed record, so the
            # extension-based media-type detection in the download
            # endpoints doesn't end up mislabeling the content.
            key = resume.s3_key if resume.s3_key.lower().endswith(".docx") else resume.s3_key.rsplit(".", 1)[0] + ".docx"
            upload_file_to_s3(
                io.BytesIO(docx_bytes), key,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            resume.s3_key = key
        except Exception as e:
            # Don't fail the save over a DOCX regen hiccup — the raw_text
            # edit itself already committed below.
            print(f"Resume DOCX regeneration on text save failed for resume {resume.id}: {e}")
            from error_logger import log_db_error
            await log_db_error(
                stage="resume_docx_regen_on_text_save",
                error=e,
                source_type="resume",
                source_id=str(resume.id),
            )
    await db.commit()
    return {"success": True, "message": "Resume text updated successfully"}

@router.get("/base/download")
async def download_base_resume(
    user_id: Optional[int] = None,
    # BUG FIX ("view keeps downloading" — same root cause and same fix as
    # phase7.py's download_application_resume): a browser can only render
    # a PDF or image inline on its own. For a .docx base resume, there is
    # no in-browser renderer at all — avoiding "attachment" headers isn't
    # enough, the browser still has nothing to show it with and falls
    # back to a download regardless. force_stream=False (View) now
    # returns a presigned Spaces URL as JSON when the file is in object
    # storage, for the frontend to hand to Google's Docs Viewer (whose
    # own servers fetch it — browser CORS never applies there).
    # force_stream=True (the actual Download button) always streams real
    # bytes through this same-origin endpoint, since a presigned URL
    # fetched directly by the browser fails — this bucket has no CORS
    # policy allowing that.
    force_stream: bool = False,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Stream the consultant's BASE resume (from their profile) on demand.

    Declared before /{id}/download so "base" is never parsed as an int id.

    Resolves consultants.base_resume_file_path at request time rather than
    embedding a presigned URL in the list response. Presigned URLs are a
    snapshot: they break when the file is replaced (old key deleted ->
    NoSuchKey) and expire after an hour. Reading fresh on each click removes
    both failure modes and needs no cache invalidation anywhere.

    Tolerates BOTH storage formats so records written before the Spaces
    migration keep working: a value that resolves to a real file on disk is
    streamed from disk, anything else is treated as a Spaces object key.
    """
    # Role scoping mirrors the base-resume injection in list_resumes().
    if current_user.role == "CONSULTANT":
        target_user_id = current_user.id
    elif current_user.role == "ADMIN":
        target_user_id = user_id or current_user.id
    elif current_user.role == "RECRUITER":
        if not user_id or user_id == current_user.id:
            target_user_id = current_user.id
        else:
            target_consultant = (await db.execute(
                select(Consultant).where(Consultant.user_id == user_id)
            )).scalar_one_or_none()
            if not target_consultant:
                raise HTTPException(status_code=404, detail="Consultant not found")
            allowed = (await db.execute(
                select(RecruiterConsultant).where(
                    RecruiterConsultant.recruiter_id == current_user.id,
                    RecruiterConsultant.consultant_id == target_consultant.id,
                    RecruiterConsultant.is_active == True,
                )
            )).scalars().first()
            if not allowed:
                raise HTTPException(status_code=403, detail="Consultant not assigned to this recruiter")
            target_user_id = user_id
    else:
        target_user_id = current_user.id

    consultant = (await db.execute(
        select(Consultant).where(Consultant.user_id == target_user_id)
    )).scalar_one_or_none()

    if not consultant or not consultant.base_resume_file_path:
        raise HTTPException(status_code=404, detail="No base resume uploaded yet.")

    stored = consultant.base_resume_file_path
    ext = os.path.splitext(stored)[1].lower()
    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext == ".docx" else "application/pdf"
    )
    display_name = f"{(consultant.full_name or 'consultant').strip().replace(' ', '_')}_base_resume{ext or '.pdf'}"

    # Legacy local-disk record (pre-Spaces migration) — no object-storage
    # key to presign, so this stays a direct stream. A .docx here has no
    # in-browser renderer and can't be handed to Google's viewer either
    # (it needs a real fetchable URL, not local bytes) — same acknowledged
    # limitation the Sent Applications endpoint (phase7.py) documents for
    # its own local-file case.
    try:
        if os.path.isfile(stored):
            with open(stored, "rb") as fh:
                body = fh.read()
            return Response(
                content=body,
                media_type=media_type,
                headers={"Content-Disposition": f'inline; filename="{display_name}"'},
            )
    except OSError:
        pass

    if not force_stream:
        # CHANGED (view now shows the actual DOCX, not a converted PDF):
        # same requirement as the Sent Applications resume-download
        # endpoint (phase7.py) — View should open the real file via
        # Google Docs Viewer (a real, publicly fetchable presigned URL,
        # fetched by Google's own servers — browser CORS never applies
        # there) instead of a server-side PDF conversion.
        presigned = generate_presigned_url(stored)
        if presigned:
            return {"url": presigned, "filename": display_name, "mimeType": media_type}

    body, content_type = await asyncio.to_thread(download_file_from_s3, stored)
    if body is None:
        raise HTTPException(
            status_code=404,
            detail="Base resume file is missing from storage. Please re-upload it in your profile.",
        )

    return Response(
        content=body,
        media_type=content_type or media_type,
        headers={"Content-Disposition": f'inline; filename="{display_name}"'},
    )

@router.get("/consultants")
async def get_consultants_for_resumes(
    # BUG FIX: the Requirements page's "Apply" link (unlike Pending
    # Applications, which already carries a specific consultantId) opens
    # ApplyToRequirementPage with no consultant context at all, so its
    # dropdown showed EVERY consultant in the system — including ones with
    # no relevance to that requirement. Optional requirement_id restricts
    # the list to consultants actually matched to that specific
    # requirement (RequirementConsultantMatch), same source the
    # Requirements table's own "Matched Consultants" column already uses.
    requirement_id: int = None,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    def map_user_consultant(u, c):
        # BUG FIX: was reading u.skills (User.skills, a JSONB column that
        # nothing in the consultant-profile flow ever writes to — profile
        # edits, resume generation, and matching all read/write
        # Consultant.primary_skills instead, see phase3.py/phase6.py/
        # resume_router.py's own resume-building code). Real consultants
        # therefore showed up with an empty skills list here, which fed
        # straight into ApplyToRequirementPage's {skills} template fill
        # and silently fell back to "relevant technologies" even when the
        # consultant had skills on file. Prefer primary_skills (+
        # secondary_skills) split into a list; keep u.skills as a
        # last-resort fallback for any user row that predates Consultant.
        consultant_skills = [
            s.strip()
            for s in ", ".join(filter(None, [
                c.primary_skills if c else None,
                c.secondary_skills if c else None,
            ])).split(",")
            if s.strip()
        ] if c else []
        return {
            "id": u.id, # Keep id for backward compatibility (maps to user_id)
            "user_id": u.id,
            "consultant_id": c.id if c else None,
            "name": u.full_name or u.email,
            "email": u.email,
            "skills": consultant_skills or u.skills,
            "experience_years": u.experience_years or (c.total_experience_years if c else 0)
        }

    matched_consultant_ids = None
    if requirement_id:
        from models import RequirementConsultantMatch, JobMatch

        # BUG FIX ("wrong candidate/email after applying from a Pending
        # Applications match"): this used to check RequirementConsultantMatch
        # only — a completely separate table from JobMatch (populated by
        # the AI matching engine in matching_router.py, which is what
        # Pending Applications actually shows matches from). Nothing ever
        # writes a JobMatch row into RequirementConsultantMatch, so a
        # consultant shown as a real match in Pending Applications could
        # be entirely absent from this list. ApplyToRequirementPage
        # resolves the consultantId passed in the URL by looking it up in
        # this exact list — if it's missing, that resolution silently
        # fails and the page falls back to a different consultant
        # (consultants?.[0]) instead, which is why the "active user"
        # shown didn't match who was actually clicked, the wrong
        # candidate got applied as, and the send went from the wrong
        # Gmail address. Union both sources so any consultant matched
        # through either system shows up correctly.
        rcm_result = await db.execute(
            select(RequirementConsultantMatch.consultant_id).where(
                RequirementConsultantMatch.requirement_id == requirement_id
            )
        )
        jobmatch_result = await db.execute(
            select(JobMatch.consultant_id).where(
                JobMatch.requirement_id == requirement_id,
                JobMatch.status != "REJECTED",
            )
        )
        matched_consultant_ids = list({
            *(row[0] for row in rcm_result.all()),
            *(row[0] for row in jobmatch_result.all()),
        })
        if not matched_consultant_ids:
            return []

    if current_user.role == "ADMIN":
        query = select(User, Consultant).join(Consultant, Consultant.user_id == User.id).where(
            User.role == "CONSULTANT",
            User.is_authorized == True,
            Consultant.status == "ACTIVE"
        )
        if matched_consultant_ids is not None:
            query = query.where(Consultant.id.in_(matched_consultant_ids))
        results = (await db.execute(query)).all()
        return [map_user_consultant(u, c) for u, c in results]
    elif current_user.role == "RECRUITER":
        consultant_users_query = select(Consultant.user_id).where(
            Consultant.status == "ACTIVE",
            or_(
                Consultant.sales_recruiter_user_id == current_user.id,
                Consultant.id.in_(
                    select(RecruiterConsultant.consultant_id).where(
                        RecruiterConsultant.recruiter_id == current_user.id
                    )
                )
            )
        )
        query = select(User, Consultant).join(Consultant, Consultant.user_id == User.id).where(
            User.id.in_(consultant_users_query),
            User.role == "CONSULTANT",
            User.is_authorized == True,
            Consultant.status == "ACTIVE"
        )
        if matched_consultant_ids is not None:
            query = query.where(Consultant.id.in_(matched_consultant_ids))
        results = (await db.execute(query)).all()
        return [map_user_consultant(u, c) for u, c in results]
    else:
        query = select(User, Consultant).join(Consultant, Consultant.user_id == User.id).where(
            User.id == current_user.id,
            User.role == "CONSULTANT",
            User.is_authorized == True,
            Consultant.status == "ACTIVE"
        )
        if matched_consultant_ids is not None:
            query = query.where(Consultant.id.in_(matched_consultant_ids))
        results = (await db.execute(query)).all()
        return [map_user_consultant(u, c) for u, c in results]

@router.get("/{id}", response_model=ResumeResponse)
async def get_resume(
    id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    # BUG FIX: was owner-only (Resume.user_id == current_user.id) — same
    # class of bug as /download. Uses the shared role-scoped helper instead.
    resume = await _get_resume_for_user(db, id, current_user)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # BACKFILL: manually-uploaded resumes only ever got data={"raw_text":
    # ...} at upload time (see /upload above) — no structured fields, so
    # the frontend routes these to the plain-text editor forever, even
    # after this same backfill was added for the base resume. Mirrors
    # get_base_resume_content's approach exactly: try the free,
    # synchronous heuristic parser first (these resumes reliably follow
    # one template, so this resolves instantly for the common case), and
    # only fall back to a real AI call when that doesn't match. raw_text
    # is kept alongside the parsed fields (not replaced) so the plain-text
    # editor and AI tailoring/matching, which both still read raw_text,
    # keep working unchanged.
    def _is_raw_text_only(d: dict) -> bool:
        return bool(d) and set(d.keys()) <= {"raw_text"}

    if _is_raw_text_only(resume.data or {}) and (resume.data.get("raw_text") or "").strip():
        raw_text = resume.data["raw_text"]
        try:
            heuristic_data = _heuristic_parse_resume_text(raw_text)
            if heuristic_data:
                resume.data = {**heuristic_data, "raw_text": raw_text}
                await db.commit()
                await db.refresh(resume)
        except Exception as e:
            print(f"Resume content backfill (heuristic) failed for resume {resume.id}: {e}")

    if _is_raw_text_only(resume.data or {}) and (resume.data.get("raw_text") or "").strip():
        raw_text = resume.data["raw_text"]
        from claude_service import parse_resume_text_to_structured_data
        try:
            parsed_data, rate_limits, usage_info = await asyncio.to_thread(
                parse_resume_text_to_structured_data, raw_text
            )
            if rate_limits:
                await save_claude_rate_limits(db, rate_limits)
            if usage_info:
                # Only persist when the AI call actually succeeded — see
                # the matching comment in get_base_resume_content for why
                # (a blank "parsing unavailable" skeleton would otherwise
                # look identical to a completed backfill and permanently
                # block every future retry).
                from phase8_ai_usage_service import log_ai_usage
                await log_ai_usage(
                    db,
                    purpose="resume_backfill",
                    model="claude-sonnet-4-6",
                    input_tokens=usage_info["input_tokens"],
                    output_tokens=usage_info["output_tokens"],
                    consultant_id=str(resume.user_id),
                )
                resume.data = {**parsed_data, "raw_text": raw_text}
                await db.commit()
                await db.refresh(resume)
            else:
                print(f"Resume content backfill returned no usage_info (AI call unavailable) for resume {resume.id}.")
        except Exception as e:
            print(f"Resume content backfill (AI) failed for resume {resume.id}: {e}")
            from error_logger import log_db_error
            await log_db_error(
                stage="resume_content_backfill",
                error=e,
                source_type="resume",
                source_id=str(resume.id),
            )

    return resume

@router.put("/{id}", response_model=ResumeResponse)
async def update_resume(
    id: int,
    request: ResumeUpdateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    # BUG FIX: was owner-only (Resume.user_id == current_user.id) — same
    # class of bug as /download. Uses the shared role-scoped helper instead.
    resume = await _get_resume_for_user(db, id, current_user)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    if request.title is not None:
        resume.title = request.title
    if request.target_role is not None:
        resume.target_role = request.target_role
    if request.job_description is not None:
        resume.job_description = request.job_description
    if request.data is not None:
        resume.data = request.data
    if request.status is not None:
        resume.status = request.status

    await db.commit()
    await db.refresh(resume)

    # BUG FIX: this endpoint only ever updated the `data` JSON column —
    # the actual downloadable PDF sitting in S3 was generated once at
    # creation time and never touched again. Editing a resume (adding a
    # LinkedIn URL, fixing a typo, or picking up a template fix like the
    # Declaration section removal) had zero effect on what "Download PDF"
    # actually served — it kept returning the exact same stale file
    # forever. Regenerate and re-upload whenever the content changed, so
    # Save Changes and Download PDF never drift apart again.
    if request.data is not None:
        from phase6 import _generate_docx, _convert_to_pdf
        from pathlib import Path

        resume_dir = Path("/tmp/resumes") / str(resume.user_id) / str(resume.id)
        resume_dir.mkdir(parents=True, exist_ok=True)
        docx_path = resume_dir / "resume.docx"
        pdf_path = resume_dir / "resume.pdf"

        try:
            _generate_docx(resume.data, docx_path)
            if _convert_to_pdf(docx_path, pdf_path):
                s3_key = f"users/{resume.user_id}/resumes/{resume.id}/resume.pdf"
                with open(pdf_path, "rb") as f:
                    if upload_file_to_s3(f, s3_key, "application/pdf"):
                        resume.s3_key = s3_key
                        resume.status = 'completed'
                        await db.commit()
                        await db.refresh(resume)
        except Exception as e:
            # Don't fail the save over a PDF regen hiccup — the data edit
            # itself already succeeded and committed above. The next save
            # (or the lazy self-heal in download_resume) will retry.
            print(f"PDF regeneration on save failed for resume {resume.id}: {e}")
            from error_logger import log_db_error
            await log_db_error(
                stage="resume_pdf_regen_on_save",
                error=e,
                source_type="resume",
                source_id=str(resume.id),
            )

    return resume

@router.delete("/{id}")
async def delete_resume(
    id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    # BUG FIX: was owner-only (Resume.user_id == current_user.id) — same
    # class of bug as /download. Uses the shared role-scoped helper instead.
    resume = await _get_resume_for_user(db, id, current_user)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    if resume.s3_key:
        delete_file_from_s3(resume.s3_key)

    await db.delete(resume)
    await db.commit()

    return {"success": True}

@router.get("/{id}/download")
async def download_resume(
    id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    # BUG FIX: was `select(Resume).where(Resume.id == id, Resume.user_id ==
    # current_user.id)` — only ever matched when the CALLER owned the resume.
    # /admin/apply/:reqId and /recruiter/apply/:reqId reuse this same page to
    # apply on behalf of a consultant, so current_user is the admin/recruiter,
    # not the consultant who owns the resume — this 404'd every single time
    # for them ("Preparing resume attachment..." -> "Failed to download and
    # attach resume"). Now uses the shared role-scoped helper (same one
    # get_resume/update_resume/delete_resume already use above) instead of
    # duplicating the ADMIN/RECRUITER/owner query logic inline: ADMIN sees
    # any resume, RECRUITER sees their own + assigned consultants', everyone
    # else only their own.
    resume = await _get_resume_for_user(db, id, current_user)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    if not resume.s3_key:
        # SELF-HEAL: some resumes (e.g. "ATS DevOps" in the bug report) have
        # tailored content (resume.data) but never got a PDF — the earlier
        # generate/finalize call hit a transient LibreOffice/S3 failure and
        # left status as 'failed_pdf_conversion'/'failed_upload' with no
        # s3_key. Previously this just 400'd forever ("Resume does not have
        # a generated PDF"), which is what the frontend's generic catch
        # turns into "Failed to download and attach resume" — the user has
        # no way to recover short of clicking "Generate Tailored Resume"
        # again from scratch. Instead, retry building the PDF from the data
        # that's already there before giving up.
        if resume.data:
            from phase6 import _generate_docx, _convert_to_pdf
            from pathlib import Path

            resume_dir = Path("/tmp/resumes") / str(resume.user_id) / str(resume.id)
            resume_dir.mkdir(parents=True, exist_ok=True)
            docx_path = resume_dir / "resume.docx"
            pdf_path = resume_dir / "resume.pdf"

            try:
                _generate_docx(resume.data, docx_path)
                if _convert_to_pdf(docx_path, pdf_path):
                    s3_key = f"users/{resume.user_id}/resumes/{resume.id}/resume.pdf"
                    with open(pdf_path, "rb") as f:
                        if upload_file_to_s3(f, s3_key, "application/pdf"):
                            resume.s3_key = s3_key
                            resume.status = 'completed'
                            await db.commit()
                            await db.refresh(resume)
            except Exception as e:
                print(f"Lazy PDF regeneration failed for resume {resume.id}: {e}")
                from error_logger import log_db_error
                await log_db_error(
                    stage="resume_lazy_pdf_regen",
                    error=e,
                    source_type="resume",
                    source_id=str(resume.id),
                )

        if not resume.s3_key:
            raise HTTPException(
                status_code=400,
                detail="This resume doesn't have a generated file yet. Click 'Generate Tailored Resume' to create one."
            )

    # CHANGED (view now shows the actual DOCX, not a converted PDF): same
    # requirement as the Sent Applications resume-download endpoint
    # (phase7.py) and the base resume's download_base_resume below — View
    # should open the real file via Google Docs Viewer (a real, publicly
    # fetchable presigned URL, which Google's own servers fetch — browser
    # CORS never applies there) instead of a server-side PDF conversion.
    #
    # FEATURE CHANGE: a generated/tailored resume's s3_key is always a
    # .pdf (see generate_resume/finalize_resume above) — the .docx it was
    # actually built from is a rendering source, not what gets stored as
    # the canonical file. Prefer showing THAT .docx here too, same as
    # manually-uploaded resumes already do, rather than the PDF that's
    # just a byproduct of it. Reuses the exact same find-or-regenerate
    # logic download_resume_docx below already has — if the .docx object
    # is missing from storage (generate_resume's own DOCX upload runs in
    # a best-effort try/except that can fail independently of the PDF
    # succeeding), rebuild it from resume.data and upload it before
    # presigning, instead of silently falling back to PDF.
    view_key = resume.s3_key
    if resume.s3_key.lower().endswith(".pdf"):
        docx_key = resume.s3_key.rsplit(".", 1)[0] + ".docx"
        docx_bytes, _ = await asyncio.to_thread(download_file_from_s3, docx_key)
        if docx_bytes is None and resume.data:
            from phase6 import _generate_docx
            from pathlib import Path

            resume_dir = Path("/tmp/resumes") / str(resume.user_id) / str(resume.id)
            resume_dir.mkdir(parents=True, exist_ok=True)
            tmp_docx_path = resume_dir / "resume_view.docx"
            try:
                _generate_docx(resume.data, tmp_docx_path)
                with open(tmp_docx_path, "rb") as f:
                    if upload_file_to_s3(f, docx_key, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                        docx_bytes = True  # just need to know it now exists
            except Exception as e:
                print(f"View-time DOCX regeneration failed for resume {resume.id}: {e}")
                from error_logger import log_db_error
                await log_db_error(
                    stage="resume_view_docx_regen",
                    error=e,
                    source_type="resume",
                    source_id=str(resume.id),
                )
        if docx_bytes is not None:
            view_key = docx_key
        # else: no .docx obtainable at all (no resume.data to rebuild
        # from) — falls through and shows the .pdf instead, same as
        # today, rather than failing View entirely.

    url = generate_presigned_url(view_key)
    if not url:
        raise HTTPException(status_code=500, detail="Failed to generate download link.")

    # Update download stats
    resume.download_count += 1
    resume.last_downloaded = datetime.now(timezone.utc)
    await db.commit()

    # BUG FIX ("generated resume view opens in drive view" instead of a
    # native PDF): this returned bare {"url": ...} with no filename or
    # mimeType. A generated/tailored resume's s3_key is always .pdf
    # already (see finalize_resume/generate_resume above) — no
    # conversion needed, this URL already points at a real PDF — but the
    # frontend's toViewableUrl has no way to know that without a mimeType
    # hint, so it defaulted to treating the file as non-viewable and
    # wrapped it in Google Docs Viewer regardless. Including mimeType
    # (derived from the real stored extension, same as download_base_resume
    # does) lets it correctly recognize an already-PDF file and skip the
    # external viewer entirely.
    import mimetypes as _mimetypes
    ext = os.path.splitext(view_key)[1].lower()
    mime_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext == ".docx"
        else _mimetypes.guess_type(view_key)[0] or "application/pdf"
    )
    safe_title = "".join(
        c for c in (resume.title or f"Resume_{id}") if c.isalnum() or c in " -_"
    ).strip() or f"Resume_{id}"

    return {"url": url, "filename": f"{safe_title}{ext or '.pdf'}", "mimeType": mime_type}

@router.get("/{id}/download/file")
async def download_resume_file(
    id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Stream the resume PDF bytes through the API.

    Why this exists alongside /{id}/download (which returns a presigned URL):
    presigned Spaces URLs work for window.open()/navigation, but a browser
    fetch() of one is a cross-origin XHR and is blocked unless the Space has
    a CORS policy. The compose/apply screens need the actual bytes to build a
    File for attachment, so they call this instead - same origin as the rest
    of the API, auth enforced, no bucket CORS required.
    """
    resume = await _get_resume_for_user(db, id, current_user)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    if not resume.s3_key:
        raise HTTPException(status_code=400, detail="Resume does not have a generated PDF.")

    body, content_type = await asyncio.to_thread(download_file_from_s3, resume.s3_key)
    if body is None:
        raise HTTPException(status_code=502, detail="Could not retrieve the resume file from storage.")

    safe_title = "".join(c for c in (resume.title or f"Resume_{id}") if c.isalnum() or c in " -_").strip()
    # BUG FIX: filename extension was hardcoded to ".pdf" regardless of the
    # actual stored file — now that uploads are DOCX-only (see /upload
    # above), this was saving downloaded files as "Title.pdf" while the
    # bytes inside were actually a .docx, which most apps then fail to open.
    # Derive it from the real stored key instead.
    ext = os.path.splitext(resume.s3_key)[1] or ".docx"
    filename = f"{safe_title or f'Resume_{id}'}{ext}"

    resume.download_count += 1
    resume.last_downloaded = datetime.now(timezone.utc)
    await db.commit()

    return Response(
        content=body,
        media_type=content_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{id}/download/file/docx")
async def download_resume_docx(
    id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """DOCX counterpart of download_resume_file above. Derives the key
    from s3_key's own naming pattern (extension swapped) rather than a
    separate column, since the upload above always writes it that way."""
    resume = await _get_resume_for_user(db, id, current_user)

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # BUG FIX ("View/Download both open the browser print dialog instead
    # of DOCX"): this used to 400 immediately with "Resume does not have
    # a generated file" whenever s3_key was empty — which is exactly the
    # case for a resume that only has structured `data` and was never
    # uploaded/finalized to an actual file. The frontend's fallback for
    # that 400 was to render the resume client-side and trigger the
    # browser's print dialog (Save as PDF), which is what was showing up
    # instead of a DOCX. Now: only attempt the S3 lookup when a s3_key
    # actually exists: no key at all skips straight to generating from
    # resume.data below, same as the "found a key but the DOCX at that
    # key is missing" self-heal case already handles.
    docx_key = resume.s3_key.rsplit(".", 1)[0] + ".docx" if resume.s3_key else None
    body, content_type = (
        await asyncio.to_thread(download_file_from_s3, docx_key) if docx_key else (None, None)
    )

    if body is None:
        # SELF-HEAL ("Resume does not have a generated DOCX" — this was the
        # error the Apply page's "Choose Resume" attach step surfaced):
        # generate/finalize_resume upload the DOCX to S3 in a try/except
        # that only logs on failure ("DOCX upload failed (PDF result
        # unaffected)") rather than failing the whole request — so a
        # resume can end up with a working PDF (resume.s3_key) but no
        # DOCX ever actually landed at docx_key. Previously that meant
        # this endpoint 400'd forever with no way to recover short of
        # regenerating the resume from scratch. Same self-heal pattern as
        # download_resume's missing-PDF case above: rebuild the DOCX from
        # resume.data (still in the DB either way) and upload it to the
        # expected key, then serve those freshly-generated bytes directly
        # instead of round-tripping back through S3.
        if resume.data:
            from phase6 import _generate_docx
            from pathlib import Path

            resume_dir = Path("/tmp/resumes") / str(resume.user_id) / str(resume.id)
            resume_dir.mkdir(parents=True, exist_ok=True)
            docx_path = resume_dir / "resume.docx"

            try:
                _generate_docx(resume.data, docx_path)
                with open(docx_path, "rb") as f:
                    docx_bytes = f.read()
                # Only upload back to S3 if there's actually a key to put
                # it at — a resume with no s3_key at all has nowhere
                # established to store it yet, so just serve the bytes
                # this once rather than inventing a storage location.
                if docx_key:
                    with open(docx_path, "rb") as f:
                        upload_file_to_s3(
                            f, docx_key,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                body = docx_bytes
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            except Exception as e:
                print(f"Lazy DOCX regeneration failed for resume {resume.id}: {e}")
                from error_logger import log_db_error
                await log_db_error(
                    stage="resume_lazy_docx_regen",
                    error=e,
                    source_type="resume",
                    source_id=str(resume.id),
                )

        if body is None:
            raise HTTPException(status_code=400, detail="Resume does not have a generated DOCX.")

    safe_title = "".join(c for c in (resume.title or f"Resume_{id}") if c.isalnum() or c in " -_").strip()
    filename = f"{safe_title or f'Resume_{id}'}.docx"

    return Response(
        content=body,
        media_type=content_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )