# phase6.py
# ---------------------------------------------------------------------------
# Phase 6 — Resume Tailoring, ATS Scoring, Resume Storage
#
# Architecture: single flat file, same pattern as phase3.py / phase4.py.
# Imports get_current_user from auth.py — no circular dependency.
#
# Endpoints:
#   POST /api/consultant/requirements/{id}/generate-resume  → GenerateResumeResultDTO
#   GET  /api/consultant/requirements/{id}/resume           → ResumeDataDTO
#   GET  /api/consultant/requirements/{id}/resume/history   → ResumeHistoryDTO
#   GET  /api/consultant/requirements/{id}/resume/download/{type}  → FileResponse (PDF/DOCX)
#   GET  /api/recruiter/consultants/{id}/requirements/{req_id}/resume → ResumeDataDTO
#
# All response shapes exactly match frontend DTO contracts in:
#   features/consultant/resume/types/index.ts
#   types/consultant.ts
# ---------------------------------------------------------------------------

from __future__ import annotations

import asyncio
import io
import json
import logging
import os
import re
import subprocess
import tempfile
import uuid
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

import httpx
from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from auth import get_current_user
from database import get_db
from claude_service import generate_tailored_resume
from models import (
    Consultant,
    ConsultantExperience,
    GeneratedResume,
    Requirement,
    RequirementConsultantMatch,
    RecruiterConsultant,
    User,
)

logger = logging.getLogger(__name__)

router = APIRouter()

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")   # legacy - unused since Claude migration
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")  # legacy - unused since Claude migration
# Kept in sync with claude_service.generate_tailored_resume()
ANTHROPIC_MODEL = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6")
ATS_PASS_THRESHOLD = 80          # per Phase 6 doc — NOT configurable
MAX_GENERATION_ATTEMPTS = 3      # per frontend RegenerateDialog MAX_ATTEMPTS = 3
RESUME_UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads/resumes"))
BASE_URL = os.getenv("API_BASE_URL", "http://localhost:8000")


# ---------------------------------------------------------------------------
# Helpers — reused patterns from phase3/phase4
# ---------------------------------------------------------------------------

def _require_role(user: User, *roles: str) -> None:
    if user.role not in roles:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail=f"Requires role: {list(roles)}",
        )


async def _get_consultant_for_user(db: AsyncSession, user: User) -> Consultant:
    result = await db.execute(select(Consultant).where(Consultant.user_id == user.id))
    consultant = result.scalars().first()
    if not consultant:
        raise HTTPException(status_code=404, detail="Consultant profile not found for this user")
    return consultant


async def _get_requirement_or_404(db: AsyncSession, requirement_id: int) -> Requirement:
    result = await db.execute(select(Requirement).where(Requirement.id == requirement_id))
    req = result.scalars().first()
    if not req:
        raise HTTPException(status_code=404, detail="Requirement not found")
    return req


async def _get_match_or_404(
    db: AsyncSession, requirement_id: int, consultant_id: int
) -> RequirementConsultantMatch:
    result = await db.execute(
        select(RequirementConsultantMatch).where(
            RequirementConsultantMatch.requirement_id == requirement_id,
            RequirementConsultantMatch.consultant_id == consultant_id,
        )
    )
    match = result.scalars().first()
    if not match:
        raise HTTPException(
            status_code=404,
            detail="No match found. Consultant must be matched to this requirement first (Phase 4).",
        )
    return match


# ---------------------------------------------------------------------------
# Task 5 — Filename convention (verbatim from Phase 6 doc code example)
# ---------------------------------------------------------------------------

def _clean_filename_part(value: Any) -> str:
    if not value:
        return "unknown"
    value = str(value).lower().strip().replace("&", "and")
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "unknown"


def _build_resume_filename(
    first_name: str,
    last_name: str,
    role: str,
    client_name: Optional[str],
    vendor_name: Optional[str],
    years_exp: Any,
) -> str:
    """Verbatim from Phase 6 doc Task 5 code example."""
    company = (
        client_name
        if client_name and client_name.upper() not in ("UNKNOWN", "N/A", "")
        else vendor_name
    )
    return (
        f"{_clean_filename_part(first_name)}_"
        f"{_clean_filename_part(last_name)}_"
        f"{_clean_filename_part(role)}_"
        f"{_clean_filename_part(company)}_"
        f"{_clean_filename_part(years_exp)}-years.pdf"
    )


# ---------------------------------------------------------------------------
# Task 3 — ATS Scoring Engine (verbatim from Phase 6 doc code example)
# ---------------------------------------------------------------------------

def _ats_score(
    jd_skills: List[str],
    resume_text: str,
    role: str,
) -> tuple[float, float, float, float, List[str], List[str]]:
    """
    Verbatim from Phase 6 doc Task 3 code example.
    Returns (total, keyword_score, role_score, format_score, matched, missing).
    """
    resume_lower = resume_text.lower()
    matched = [s for s in jd_skills if s.lower() in resume_lower]
    missing = [s for s in jd_skills if s.lower() not in resume_lower]
    keyword_score = len(matched) / max(len(jd_skills), 1) * 70
    role_score = 15.0 if role.lower() in resume_lower else 5.0
    format_score = 15.0
    total = round(min(100.0, keyword_score + role_score + format_score), 2)
    return (
        total,
        round(keyword_score, 2),
        round(role_score, 2),
        round(format_score, 2),
        matched,
        missing,
    )


# ---------------------------------------------------------------------------
# Task 2 — AI Resume Tailoring System Prompt (verbatim from Phase 6 doc)
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are a resume tailoring assistant. Use only the consultant base resume and structured experience provided.
Do not invent clients, projects, skills, certifications, titles, dates, or years of experience.
If the job description includes a skill not present in the consultant profile, mark it as missing instead of adding it.
Improve wording, reorder relevant skills, and add 4-6 truthful bullets based on existing experience only.
Return structured JSON.

Return exactly this JSON structure with no markdown, no code fences, no extra text:
{
  "name": "string",
  "email": "string",
  "phone": "string",
  "summary": "string",
  "skills": ["skill1", "skill2"],
  "missing_skills": ["skill_not_in_profile"],
  "experience": [
    {
      "client": "string",
      "role": "string",
      "start": "string",
      "end": "string",
      "location": "string",
      "bullets": ["bullet1", "bullet2", "bullet3", "bullet4"]
    }
  ],
  "generation_notes": "string"
}"""


async def _call_ai_tailoring(
    consultant: Consultant,
    experiences: List[ConsultantExperience],
    requirement: Requirement,
    matched_skills: List[str],
    missing_skills: List[str],
    db: Optional[AsyncSession] = None,
) -> dict:
    """
    Generate a tailored resume via Anthropic Claude (claude_service).

    MIGRATION NOTE: this previously called OpenAI/GPT-4o directly. The rest of
    the platform (resume_router -> claude_service) standardised on Claude and
    OPENAI_API_KEY was never provisioned in any environment, so this code path
    failed for every consultant while "My Resumes" worked. Same JSON contract,
    one provider, one key to manage.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key or api_key.startswith("your_"):
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="ANTHROPIC_API_KEY is not configured. Set it in .env to enable resume generation.",
        )

    # ── Structured profile ────────────────────────────────────────────────
    # Key names mirror claude_service's internal mock_fallback shape so its
    # offline fallback still yields sensible data instead of placeholders.
    experience_payload: List[Dict[str, Any]] = []
    for exp in sorted(experiences, key=lambda e: e.sort_order):
        raw = f"{exp.responsibilities or ''}\n{exp.achievements or ''}"
        bullets = [ln.strip(" -*\u2022\t") for ln in re.split(r"[\r\n]+", raw) if ln.strip(" -*\u2022\t")]
        experience_payload.append(
            {
                "company": exp.client_name or "",
                "role": exp.role_title or "",
                "start_date": exp.start_date.strftime("%b %Y") if exp.start_date else "",
                "end_date": "Present"
                if exp.is_present
                else (exp.end_date.strftime("%b %Y") if exp.end_date else ""),
                "technologies": exp.technologies or [],
                "bullets": bullets,
            }
        )

    primary = [s.strip() for s in (consultant.primary_skills or "").split(",") if s.strip()]
    secondary = [s.strip() for s in (consultant.secondary_skills or "").split(",") if s.strip()]

    # BUG FIX: this resume_info was built ONLY from primary_skills/
    # secondary_skills (comma-separated text columns) and ConsultantExperience
    # table rows — completely separate from consultant.resume_info, the JSON
    # blob the "My Resumes" flow (resume_router.py) reads from. A consultant
    # whose real profile data was entered/imported into that JSON blob (e.g.
    # via the resume JSON import feature) but never duplicated into these
    # separate structured fields would get an almost entirely empty
    # resume_info here — no skills, no experience, no education, no
    # certifications at all (education/certifications weren't sourced from
    # anywhere in this function, structured or JSON) — which is exactly the
    # shape of failure that produces a "resume is basically empty" result.
    # Fall back to the JSON blob for anything the structured fields didn't
    # provide, so both data sources actually work.
    profile_json = consultant.resume_info or {}
    tech_stack_json = profile_json.get("tech_stack") or {}

    if not primary and not secondary:
        primary = tech_stack_json.get("expert") or profile_json.get("skills") or []
        secondary = (
            (tech_stack_json.get("exposure") or [])
            + (tech_stack_json.get("familiar") or [])
        )

    if not experience_payload:
        experience_payload = [
            {
                "company": exp.get("company", ""),
                "role": exp.get("role", ""),
                "start_date": exp.get("start_date", ""),
                "end_date": exp.get("end_date", "Present"),
                "technologies": exp.get("tech_used", []),
                "bullets": exp.get("bullets", []),
            }
            for exp in profile_json.get("experience", [])
        ]

    resume_info: Dict[str, Any] = {
        "full_name": consultant.full_name or profile_json.get("full_name", ""),
        "email": consultant.email or profile_json.get("email", ""),
        "phone": consultant.phone or profile_json.get("phone", ""),
        "total_experience_years": float(consultant.total_experience_years)
        if consultant.total_experience_years is not None
        else profile_json.get("years_experience"),
        "work_authorization": consultant.work_authorization or profile_json.get("visa_type", ""),
        "current_location": consultant.current_location or profile_json.get("location", ""),
        "linkedin": profile_json.get("linkedin", ""),
        "github": profile_json.get("github", ""),
        "tech_stack": {"expert": primary, "familiar": secondary},
        "base_resume_text": consultant.base_resume_text or "",
        "experience": experience_payload,
        "education": profile_json.get("education") or profile_json.get("educational_background") or [],
        "certifications": profile_json.get("certifications") or [],
    }

    # ── Requirement context (preserves everything the old prompt carried) ──
    jd_context = f"""Role: {requirement.role}
Client: {requirement.client or 'Not specified'}
Location: {requirement.location or 'Not specified'}
Work Mode: {requirement.work_mode or 'Not specified'}
Employment Types: {', '.join(requirement.employment_types or []) or 'Not specified'}

JOB DESCRIPTION:
{requirement.job_description or 'No job description available.'}

MATCHED SKILLS (already in profile): {', '.join(matched_skills) or 'None'}
MISSING SKILLS (in JD, not in profile): {', '.join(missing_skills) or 'None'}"""

    # ── Call Claude off the event loop ────────────────────────────────────
    # generate_tailored_resume() is synchronous and blocks for the duration of
    # the HTTP call. Running it inline would freeze this uvicorn worker for
    # every other request, so it goes to a thread.
    #
    # BUG FIX: generate_tailored_resume() returns a 3-tuple
    # (result_json, rate_limits, usage_info) on every code path — success,
    # no-API-key, and the exception fallback all return three values. This
    # was only unpacking two, so EVERY call raised
    # "ValueError: too many values to unpack (expected 2)", caught by the
    # except block below and surfaced as a generic "AI service error"
    # regardless of whether generation would have actually succeeded.
    try:
        resume_data, rate_limits, usage_info = await asyncio.to_thread(
            generate_tailored_resume, resume_info, jd_context
        )
    except Exception as exc:  # noqa: BLE001 - surface any client/transport error
        logger.error("Claude resume generation failed: %s", exc)
        raise HTTPException(status_code=502, detail=f"AI service error: {exc}")

    if not isinstance(resume_data, dict) or not resume_data.get("experience"):
        logger.error("Claude returned unusable payload: %r", resume_data)
        raise HTTPException(status_code=502, detail="AI returned malformed response. Please retry.")

    # claude_service swallows API errors and silently returns mock data.
    # Persisting that as a real resume would be worse than failing loudly.
    if "Mock generated" in (resume_data.get("generation_notes") or ""):
        logger.error(
            "Claude returned mock fallback (consultant_id=%s requirement_id=%s) - check ANTHROPIC_API_KEY",
            consultant.id, requirement.id,
        )
        raise HTTPException(
            status_code=502,
            detail="AI service is unavailable right now. Please retry in a moment.",
        )

    # Best-effort telemetry — powers the admin AI Usage screen. Never fatal.
    if db is not None and rate_limits:
        try:
            from phase8_ai_usage_service import save_claude_rate_limits

            await save_claude_rate_limits(db, rate_limits)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to persist Claude rate limits: %s", exc)

    return resume_data


def _validate_resume_output(resume_data: dict, consultant: Consultant) -> tuple[dict, str]:
    """
    Task 2: Reject skills not in consultant profile.
    Returns (validated_data, generation_notes).
    """
    notes = resume_data.get("generation_notes", "")
    profile_skills_lower = set(
        s.strip().lower()
        for s in (consultant.primary_skills or "").split(",") + (consultant.secondary_skills or "").split(",")
        if s.strip()
    )

    original_skills = resume_data.get("skills", [])
    validated_skills = []
    rejected_skills = []

    for skill in original_skills:
        if skill.strip().lower() in profile_skills_lower:
            validated_skills.append(skill)
        else:
            rejected_skills.append(skill)

    if rejected_skills:
        note = f"Rejected unsupported skills: {', '.join(rejected_skills)}."
        notes = f"{notes} {note}".strip()
        logger.warning(
            "Rejected %d invented skills for consultant_id=%s: %s",
            len(rejected_skills), consultant.id, rejected_skills,
        )

    resume_data["skills"] = validated_skills
    resume_data["generation_notes"] = notes
    return resume_data, notes


# ---------------------------------------------------------------------------
# Task 4 — DOCX Generation (verbatim from Phase 6 doc code example)
# ---------------------------------------------------------------------------

def _generate_docx(resume_data: dict, output_path: Path) -> None:
    """Master Resume DOCX Builder modeled after shivashankar.docx.pdf template."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import re

    doc = Document()

    # Set 0.75-inch page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def add_section_header(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.underline = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_formatted_paragraph(text: str, style: Optional[str] = None, space_after: int = 4):
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(space_after)
        if not text:
            return p
        normalized = re.sub(r'</?(b|strong)>', '**', str(text))
        parts = re.split(r'(\*\*.*?\*\*)', normalized)
        for part in parts:
            if part.startswith("**") and part.endswith("**") and len(part) > 4:
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        return p

    # 1. HEADER & CONTACT
    name = resume_data.get("name", "").strip() or "FULL NAME"
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(15)

    contact_parts = []
    if resume_data.get("email"):
        contact_parts.append(resume_data["email"])
    if resume_data.get("phone"):
        contact_parts.append(f"Mobile. No: {resume_data['phone']}")
    if resume_data.get("location"):
        contact_parts.append(resume_data["location"])
    if resume_data.get("linkedin"):
        contact_parts.append(resume_data["linkedin"])
    if resume_data.get("github"):
        contact_parts.append(resume_data["github"])

    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.paragraph_format.space_after = Pt(6)
        p_contact.add_run(" | ".join(contact_parts))

    # Add thin horizontal rule under header
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(8)
    p_hr_border = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')
    p_hr_border.append(bottom_border)
    p_hr._p.get_or_add_pPr().append(p_hr_border)

    # 2. CAREER OBJECTIVE / SUMMARY
    career_obj = resume_data.get("career_objective") or resume_data.get("summary")
    if career_obj:
        add_section_header("CAREER OBJECTIVE:")
        add_formatted_paragraph(career_obj, space_after=6)

    # 3. TECHNICAL PROFICIENCIES — moved here, right after Career
    # Objective and before Experience, to match the web preview/editor's
    # section order. Rendered as a real bordered table (was a
    # colon-padded bullet list before, which is why it looked nothing
    # like the bordered category/skills table shown in the preview).
    tech_profs = resume_data.get("technical_proficiencies", [])
    skills_list = resume_data.get("skills", [])
    if tech_profs or skills_list:
        add_section_header("TECHNICAL PROFICIENCIES:")
        if tech_profs and isinstance(tech_profs, list):
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            table.autofit = False
            table.columns[0].width = Inches(1.9)
            table.columns[1].width = Inches(4.6)
            for tp in tech_profs:
                cat = tp.get("category", "Skills")
                skills_val = ", ".join(tp.get("skills", [])) if isinstance(tp.get("skills"), list) else str(tp.get("skills", ""))
                row = table.add_row()
                cell_cat, cell_skills = row.cells
                cell_cat.width = Inches(1.9)
                cell_skills.width = Inches(4.6)
                r_cat = cell_cat.paragraphs[0].add_run(cat)
                r_cat.bold = True
                r_cat.font.size = Pt(9)
                r_sk = cell_skills.paragraphs[0].add_run(skills_val)
                r_sk.font.size = Pt(9)
            # Spacer after the table so the next section header isn't
            # crammed against it.
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)
        elif skills_list:
            p_sk = doc.add_paragraph(style="List Bullet")
            r_cat = p_sk.add_run(f"{'Core Skills':<25} : ")
            r_cat.bold = True
            p_sk.add_run(", ".join(skills_list))

    # 4. EXPERIENCE
    experience = resume_data.get("experience", [])
    if experience:
        add_section_header("EXPERIENCE:")
        for exp in experience:
            company = exp.get("client") or exp.get("company") or ""
            role = exp.get("role") or exp.get("title") or ""
            start = exp.get("start") or exp.get("start_date") or ""
            end = exp.get("end") or exp.get("end_date") or ""
            date_str = f"{start} – {end}".strip(" –")

            exp_line = f"Associated with {company}" if company else ""
            if date_str:
                exp_line += f" ({date_str})"
            if exp_line:
                p_exp = add_formatted_paragraph(exp_line)
                p_exp.runs[0].bold = True

            if role:
                add_formatted_paragraph(f"Designation: {role}")
            if exp.get("location"):
                add_formatted_paragraph(f"Location: {exp['location']}")
            if exp.get("description"):
                add_formatted_paragraph(f"Job Description: {exp['description']}")

            for bullet in exp.get("bullets", []):
                add_formatted_paragraph(bullet, style="List Bullet")

    # 5. KEY PROJECTS
    key_projects = resume_data.get("key_projects", [])
    if key_projects:
        add_section_header("KEY PROJECTS:")
        for idx, proj in enumerate(key_projects, 1):
            title = proj.get("title") or proj.get("game_name") or f"Project {idx}"
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_after = Pt(2)
            r_t = p_proj.add_run(f"{idx}) Project Title: {title}")
            r_t.bold = True

            if proj.get("description"):
                add_formatted_paragraph(f"Description: {proj['description']}")
            if proj.get("role"):
                add_formatted_paragraph(f"Role: {proj['role']}")

            resps = proj.get("responsibilities", [])
            if resps:
                add_formatted_paragraph("Responsibilities:")
                for resp in resps:
                    add_formatted_paragraph(resp, style="List Bullet")

            if proj.get("team_size"):
                add_formatted_paragraph(f"Team Size: {proj['team_size']}")
            if proj.get("duration"):
                add_formatted_paragraph(f"Duration: {proj['duration']}")
            if proj.get("technical_tools") or proj.get("tools"):
                tools = proj.get("technical_tools") or proj.get("tools")
                tools_str = ", ".join(tools) if isinstance(tools, list) else str(tools)
                add_formatted_paragraph(f"Technical Tools: {tools_str}")

    # 6. OTHER PROJECTS
    other_projects = resume_data.get("other_projects")
    if other_projects:
        add_section_header("OTHER PROJECTS:")
        if isinstance(other_projects, str):
            add_formatted_paragraph(other_projects)
        elif isinstance(other_projects, list):
            for op in other_projects:
                if isinstance(op, str):
                    add_formatted_paragraph(op, style="List Bullet")
                elif isinstance(op, dict):
                    add_formatted_paragraph(f"{op.get('title', '')}: {op.get('description', '')}")

    # 7. ACADEMIC PROJECTS
    academic_projects = resume_data.get("academic_projects", [])
    if academic_projects:
        add_section_header("ACADEMIC PROJECTS:")
        for idx, proj in enumerate(academic_projects, 1):
            title = proj.get("title") or f"Project {idx}"
            p_proj = doc.add_paragraph()
            p_proj.paragraph_format.space_after = Pt(2)
            r_t = p_proj.add_run(f"{idx}) Project Title: {title}")
            r_t.bold = True

            if proj.get("platform"):
                add_formatted_paragraph(f"Platform: {proj['platform']}")
            if proj.get("description"):
                add_formatted_paragraph(f"Description: {proj['description']}")
            if proj.get("role"):
                add_formatted_paragraph(f"Role: {proj['role']}")

            resps = proj.get("responsibilities", [])
            if resps:
                for resp in resps:
                    add_formatted_paragraph(resp, style="List Bullet")

            if proj.get("team_size"):
                add_formatted_paragraph(f"Team Size: {proj['team_size']}")
            if proj.get("duration"):
                add_formatted_paragraph(f"Duration: {proj['duration']}")
            if proj.get("technical_tools"):
                tools_str = ", ".join(proj["technical_tools"]) if isinstance(proj["technical_tools"], list) else str(proj["technical_tools"])
                add_formatted_paragraph(f"Technical Tools: {tools_str}")

    # 8. EDUCATIONAL BACKGROUND
    education = resume_data.get("education") or resume_data.get("educational_background")
    if education:
        add_section_header("EDUCATIONAL BACKGROUND:")
        if isinstance(education, list):
            for edu in education:
                if isinstance(edu, dict):
                    deg = edu.get("degree", "")
                    inst = edu.get("institution") or edu.get("college", "")
                    yr = edu.get("year", "")
                    det = edu.get("details") or edu.get("percentage", "")
                    line = f"{deg} - {inst} ({yr})" if inst else deg
                    if det:
                        line += f" | {det}"
                    add_formatted_paragraph(line, style="List Bullet")
                else:
                    add_formatted_paragraph(str(edu), style="List Bullet")
        elif isinstance(education, str):
            add_formatted_paragraph(education)

    # 9. NON-TECHNICAL PROFICIENCIES
    non_tech = resume_data.get("non_technical_proficiencies")
    if non_tech:
        add_section_header("NON-TECHNICAL PROFICIENCIES:")
        if isinstance(non_tech, list):
            for nt in non_tech:
                add_formatted_paragraph(str(nt), style="List Bullet")
        else:
            add_formatted_paragraph(str(non_tech))

    # 10. ACHIEVEMENTS
    achievements = resume_data.get("achievements")
    if achievements:
        add_section_header("ACHIEVEMENTS:")
        if isinstance(achievements, list):
            for ach in achievements:
                add_formatted_paragraph(str(ach), style="List Bullet")
        else:
            add_formatted_paragraph(str(achievements))

    # 11. HOBBIES & INTEREST
    hobbies = resume_data.get("hobbies_and_interests") or resume_data.get("hobbies")
    if hobbies:
        add_section_header("HOBBIES & INTEREST:")
        if isinstance(hobbies, list):
            for h in hobbies:
                add_formatted_paragraph(str(h), style="List Bullet")
        else:
            add_formatted_paragraph(str(hobbies))

    # Missing skills transparency
    missing = resume_data.get("missing_skills", [])
    if missing:
        add_section_header("SKILLS GAP (NOT IN PROFILE):")
        add_formatted_paragraph(f"Skills requested in job description: {', '.join(missing)}")

    doc.save(str(output_path))


def _convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """
    LibreOffice headless first (production), reportlab fallback (local dev).
    Returns True if PDF was created successfully.
    """
    # Try LibreOffice headless
    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(pdf_path.parent), str(docx_path)],
            capture_output=True, timeout=30, text=True,
        )
        if result.returncode == 0:
            lo_output = pdf_path.parent / (docx_path.stem + ".pdf")
            if lo_output.exists() and lo_output != pdf_path:
                lo_output.rename(pdf_path)
            elif lo_output.exists():
                pass  # already at correct path
            return pdf_path.exists()
        logger.warning("LibreOffice failed: %s", result.stderr[:200])
    except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
        logger.warning("LibreOffice not available (%s), using reportlab fallback.", exc)

    # reportlab fallback
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from docx import Document as DocxDocument

        # Extract text from DOCX for the PDF
        try:
            docx_doc = DocxDocument(str(docx_path))
            paragraphs_text = [p.text for p in docx_doc.paragraphs if p.text.strip()]
        except Exception:
            paragraphs_text = ["Resume content — install LibreOffice for formatted PDF."]

        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for text in paragraphs_text[:50]:  # limit for safety
            story.append(Paragraph(text, styles["Normal"]))
            story.append(Spacer(1, 6))
        pdf_doc.build(story)
        return True
    except ImportError:
        logger.error("Neither LibreOffice nor reportlab available for PDF.")
        return False


# ---------------------------------------------------------------------------
# Download URL helpers
# ---------------------------------------------------------------------------

def _build_download_url(requirement_id: int, file_type: str) -> str:
    """Build the download endpoint URL for a file type (pdf or docx)."""
    return f"{BASE_URL}/api/consultant/requirements/{requirement_id}/resume/download/{file_type}"


def _build_download_urls(requirement_id: int) -> dict:
    """
    Build fresh download URLs.
    Frontend spec: 'fresh presigned URLs generated on every fetchResumeData call'
    and 'staleTime: 0 — NO URL caching allowed per spec'.
    URLs expire in 24h (displayed in DownloadButtonGroup component).
    """
    expires_at = (datetime.now(timezone.utc) + timedelta(hours=24)).isoformat()
    return {
        "pdfUrl": _build_download_url(requirement_id, "pdf"),
        "docxUrl": _build_download_url(requirement_id, "docx"),
        "expiresAt": expires_at,
    }


# ---------------------------------------------------------------------------
# Response builders — exactly matching frontend DTO contracts
# ---------------------------------------------------------------------------

def _build_resume_data_dto(
    generated: GeneratedResume,
    requirement: Requirement,
) -> dict:
    """
    Build ResumeDataDTO matching frontend types/index.ts:
    { requirementId, requirementRole, clientName, jdText,
      atsScore, atsBreakdown, skillMatch, downloadUrls,
      generationAttempts, generated }
    """
    return {
        "requirementId": str(requirement.id),
        "requirementRole": requirement.role,
        "clientName": requirement.client or "",
        "jdText": requirement.job_description or "",
        "atsScore": float(generated.ats_score or 0),
        "atsBreakdown": {
            "keywordMatch": float(generated.ats_keyword_score or 0),
            "roleTitleMatch": float(generated.ats_role_score or 0),
            "formatScore": float(generated.ats_format_score or 0),
        },
        "skillMatch": {
            "matched": generated.ats_matched_keywords or [],
            "missing": generated.ats_missing_keywords or [],
        },
        "downloadUrls": _build_download_urls(requirement.id),
        "generationAttempts": generated.generation_attempt,
        "generated": True,
    }


def _build_generate_result_dto(
    generated: GeneratedResume,
    requirement_id: int,
) -> dict:
    """
    Build GenerateResumeResultDTO matching frontend types/index.ts:
    { requirementId, atsScore, atsBreakdown, skillMatch,
      downloadUrls, generationAttempts }
    """
    return {
        "requirementId": str(requirement_id),
        "atsScore": float(generated.ats_score or 0),
        "atsBreakdown": {
            "keywordMatch": float(generated.ats_keyword_score or 0),
            "roleTitleMatch": float(generated.ats_role_score or 0),
            "formatScore": float(generated.ats_format_score or 0),
        },
        "skillMatch": {
            "matched": generated.ats_matched_keywords or [],
            "missing": generated.ats_missing_keywords or [],
        },
        "downloadUrls": _build_download_urls(requirement_id),
        "generationAttempts": generated.generation_attempt,
    }


# ---------------------------------------------------------------------------
# Core generation pipeline
# ---------------------------------------------------------------------------

async def _run_generation_pipeline(
    db: AsyncSession,
    consultant: Consultant,
    requirement: Requirement,
    match: RequirementConsultantMatch,
    current_user: User,
    attempt: int = 1,
) -> GeneratedResume:
    """
    Full Phase 6 pipeline per doc flow:
    AI Tailor → Validate → ATS Score → if <80 retry once → DOCX → PDF → store
    Max attempts: 3 (matching frontend MAX_ATTEMPTS = 3).
    """
    # Load experiences (batch — no N+1)
    exp_result = await db.execute(
        select(ConsultantExperience)
        .where(ConsultantExperience.consultant_id == consultant.id)
        .order_by(ConsultantExperience.sort_order.asc())
    )
    experiences = exp_result.scalars().all()

    matched_skills = match.matched_skills or []
    missing_skills = match.missing_skills or []

    logger.info(
        "AI resume generation: consultant_id=%s requirement_id=%s attempt=%d",
        consultant.id, requirement.id, attempt,
    )

    # ── Step 1: Call AI ───────────────────────────────────────────────────
    resume_data = await _call_ai_tailoring(
        consultant, experiences, requirement, matched_skills, missing_skills, db
    )

    # ── Step 2: Validate — reject invented skills ─────────────────────────
    resume_data, generation_notes = _validate_resume_output(resume_data, consultant)

    # ── Step 3: Build resume text for ATS scoring ─────────────────────────
    resume_text_parts = [
        resume_data.get("summary", ""),
        " ".join(resume_data.get("skills", [])),
        requirement.role,
    ]
    for exp in resume_data.get("experience", []):
        resume_text_parts.append(exp.get("role", ""))
        resume_text_parts.append(" ".join(exp.get("bullets", [])))
    resume_text = " ".join(resume_text_parts)

    # ── Step 4: ATS score ─────────────────────────────────────────────────
    jd_skills = matched_skills + missing_skills
    ats_total, kw_score, role_score, fmt_score, ats_matched, ats_missing = _ats_score(
        jd_skills, resume_text, requirement.role
    )
    logger.info("ATS score=%s attempt=%d consultant_id=%s", ats_total, attempt, consultant.id)

    # ── Step 5: One retry if below threshold (per doc Task 3) ─────────────
    if ats_total < ATS_PASS_THRESHOLD and attempt < MAX_GENERATION_ATTEMPTS:
        logger.info("ATS %s < %s — retrying (attempt %d)", ats_total, ATS_PASS_THRESHOLD, attempt + 1)
        return await _run_generation_pipeline(
            db, consultant, requirement, match, current_user, attempt=attempt + 1
        )

    final_status = "READY" if ats_total >= ATS_PASS_THRESHOLD else "NEEDS_REVIEW"
    if final_status == "NEEDS_REVIEW":
        logger.warning(
            "ATS %s < %s after %d attempts — NEEDS_REVIEW consultant_id=%s",
            ats_total, ATS_PASS_THRESHOLD, attempt, consultant.id,
        )

    # ── Step 6: Build filename (Task 5) ──────────────────────────────────
    full_name = (consultant.full_name or "").strip()
    name_parts = full_name.split() if full_name else ["", ""]
    first_name = name_parts[0] if name_parts else ""
    last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

    base_filename_with_ext = _build_resume_filename(
        first_name=first_name,
        last_name=last_name,
        role=requirement.role,
        client_name=requirement.client,
        vendor_name=requirement.vendor,
        years_exp=consultant.total_experience_years,
    )
    # base_filename_with_ext ends in .pdf per doc — strip for DOCX variant
    base_stem = base_filename_with_ext.removesuffix(".pdf")

    # ── Step 7: Generate DOCX ─────────────────────────────────────────────
    resume_dir = RESUME_UPLOAD_DIR / "generated" / str(consultant.id) / str(requirement.id)
    resume_dir.mkdir(parents=True, exist_ok=True)

    docx_path = resume_dir / f"{base_stem}.docx"
    pdf_path = resume_dir / f"{base_stem}.pdf"

    try:
        _generate_docx(resume_data, docx_path)
        logger.info("DOCX generated: %s", docx_path)
    except Exception as exc:
        logger.error("DOCX generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX: {exc}")

    # ── Step 8: Convert to PDF ────────────────────────────────────────────
    pdf_ok = _convert_to_pdf(docx_path, pdf_path)
    s3_pdf_key = None
    if not pdf_ok:
        logger.warning("PDF conversion failed — DOCX available, PDF unavailable")
    else:
        try:
            from s3_service import upload_file_to_s3
            import uuid
            s3_pdf_key = f"resumes/generated/{consultant.id}/{requirement.id}/{uuid.uuid4()}.pdf"
            with open(pdf_path, "rb") as f:
                upload_success = upload_file_to_s3(f, s3_pdf_key, "application/pdf")
            if upload_success:
                logger.info(f"Uploaded generated PDF to S3: {s3_pdf_key}")
                # Optional: Delete local PDF to save space
                try:
                    pdf_path.unlink(missing_ok=True)
                except Exception as del_err:
                    logger.warning(f"Could not delete local PDF: {del_err}")
            else:
                logger.warning("S3 upload returned False, keeping local PDF")
                s3_pdf_key = None
        except Exception as e:
            logger.error(f"S3 upload failed for generated PDF: {e}")
            s3_pdf_key = None

    # ── Step 9: Mark previous versions non-final ─────────────────────────
    await db.execute(
        update(GeneratedResume)
        .where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement.id,
            GeneratedResume.is_final == True,
        )
        .values(is_final=False)
    )

    # ── Step 10: Save generated_resume record ────────────────────────────
    generated = GeneratedResume(
        consultant_id=consultant.id,
        requirement_id=requirement.id,
        created_by_user_id=current_user.id,
        ai_model=ANTHROPIC_MODEL,
        generation_notes=generation_notes,
        generation_attempt=attempt,
        resume_content=resume_data,
        ats_score=ats_total,
        ats_keyword_score=kw_score,
        ats_role_score=role_score,
        ats_format_score=fmt_score,
        ats_matched_keywords=ats_matched,
        ats_missing_keywords=ats_missing,
        docx_path=str(docx_path),
        pdf_path=s3_pdf_key if s3_pdf_key else (str(pdf_path) if pdf_ok else None),
        pdf_url=f"/api/consultant/requirements/{requirement.id}/resume/download/pdf" if pdf_ok else None,
        filename=base_filename_with_ext,
        status=final_status,
        generation_status="COMPLETED" if final_status == "READY" else final_status,
        is_final=True,
    )
    db.add(generated)

    # ── Step 11: Update match status ─────────────────────────────────────
    match.status = "READY_TO_APPLY" if final_status == "READY" else "RESUME_GENERATED"

    await db.commit()
    await db.refresh(generated)

    logger.info(
        "Generation complete: id=%s ats=%s status=%s match_status=%s",
        generated.id, ats_total, final_status, match.status,
    )
    return generated


# ---------------------------------------------------------------------------
# Task 1 — Generate Resume API
# ---------------------------------------------------------------------------

@router.post(
    "/api/consultant/requirements/{requirement_id}/generate-resume",
    summary="Generate AI-tailored resume (Task 1) — returns GenerateResumeResultDTO",
)
async def generate_resume(
    requirement_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    CONSULTANT role only — per doc code example: require_role('CONSULTANT').
    Returns GenerateResumeResultDTO matching frontend types/index.ts.
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)
    requirement = await _get_requirement_or_404(db, requirement_id)
    match = await _get_match_or_404(db, requirement_id, consultant.id)

    # Check max attempts — frontend MAX_ATTEMPTS = 3
    existing_count_result = await db.execute(
        select(GeneratedResume).where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement_id,
        )
    )
    existing_count = len(existing_count_result.scalars().all())
    if existing_count >= MAX_GENERATION_ATTEMPTS:
        raise HTTPException(
            status_code=422,
            detail=f"Maximum {MAX_GENERATION_ATTEMPTS} generation attempts reached for this requirement.",
        )

    # Validate prerequisites
    if not consultant.base_resume_text and not consultant.primary_skills:
        raise HTTPException(
            status_code=422,
            detail="Upload a base resume or add skills before generating a tailored resume.",
        )
    if not requirement.job_description:
        raise HTTPException(
            status_code=422,
            detail="This requirement has no job description. Cannot generate a tailored resume.",
        )

    generated = await _run_generation_pipeline(
        db=db,
        consultant=consultant,
        requirement=requirement,
        match=match,
        current_user=current_user,
        attempt=existing_count + 1,
    )

    try:
        from phase5 import _broadcast_event
        await _broadcast_event("resume_generated", {
            "resume_id": str(generated.id),
            "consultant_id": str(consultant.id),
            "requirement_id": str(requirement_id),
            "status": generated.generation_status,
            "ats_score": float(generated.ats_score) if generated.ats_score is not None else None,
        })
    except ImportError:
        pass

    return _build_generate_result_dto(generated, requirement_id)


# ---------------------------------------------------------------------------
# Task 6 — Resume Preview APIs (GET endpoints)
# ---------------------------------------------------------------------------

@router.get(
    "/api/consultant/requirements/{requirement_id}/resume",
    summary="Get current resume data — returns ResumeDataDTO (fresh download URLs every call)",
)
async def get_resume(
    requirement_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Returns ResumeDataDTO exactly matching frontend types/index.ts.
    staleTime: 0 in frontend — fresh download URLs on every call per spec.
    Returns null-equivalent (404) if no resume exists → frontend redirects to dashboard.
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)
    requirement = await _get_requirement_or_404(db, requirement_id)

    result = await db.execute(
        select(GeneratedResume).where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement_id,
            GeneratedResume.is_final == True,
        )
    )
    generated = result.scalars().first()
    if not generated:
        raise HTTPException(status_code=404, detail="No generated resume found.")

    return _build_resume_data_dto(generated, requirement)


class GeneratedResumeContentDTO(BaseModel):
    resumeContent: dict
    requirementRole: str
    clientName: str


@router.get(
    "/api/consultant/requirements/{requirement_id}/resume/content",
    response_model=GeneratedResumeContentDTO,
    summary="Get the editable structured content behind a tailored resume",
)
async def get_resume_content(
    requirement_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    NEW: powers the dashboard's Edit action for tailored resumes. Nothing
    previously exposed resume_content — only the PDF/DOCX and ATS
    breakdown, via get_resume above. Same authorization lookup as that
    endpoint, just returning the structured JSON instead of file links.
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)
    requirement = await _get_requirement_or_404(db, requirement_id)

    result = await db.execute(
        select(GeneratedResume).where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement_id,
            GeneratedResume.is_final == True,
        )
    )
    generated = result.scalars().first()
    if not generated:
        raise HTTPException(status_code=404, detail="No generated resume found.")

    return GeneratedResumeContentDTO(
        resumeContent=generated.resume_content or {},
        requirementRole=requirement.role,
        clientName=requirement.client or "",
    )


class UpdateGeneratedResumeRequest(BaseModel):
    resumeContent: dict


@router.put(
    "/api/consultant/requirements/{requirement_id}/resume/content",
    summary="Save edits to a tailored resume and regenerate its PDF/DOCX",
)
async def update_resume_content(
    requirement_id: int,
    request: UpdateGeneratedResumeRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    NEW: the dashboard previously had no way to edit a tailored resume at
    all — this saves the edited content AND regenerates the actual
    downloadable files, so the PDF/DOCX a consultant downloads always
    matches what they last edited. Re-scores ATS against the same
    requirement/match so the score stays meaningful after edits, using
    the same pipeline as initial generation (_generate_docx,
    _convert_to_pdf, upload_file_to_s3).
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)
    requirement = await _get_requirement_or_404(db, requirement_id)

    result = await db.execute(
        select(GeneratedResume).where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement_id,
            GeneratedResume.is_final == True,
        )
    )
    generated = result.scalars().first()
    if not generated:
        raise HTTPException(status_code=404, detail="No generated resume found.")

    resume_data = request.resumeContent
    generated.resume_content = resume_data

    # Re-score against the same requirement/match this resume was
    # originally tailored for, so editing doesn't silently zero the score.
    match_result = await db.execute(
        select(RequirementConsultantMatch).where(
            RequirementConsultantMatch.consultant_id == consultant.id,
            RequirementConsultantMatch.requirement_id == requirement_id,
        )
    )
    match = match_result.scalars().first()
    matched_skills = (match.matched_skills if match else None) or []
    missing_skills = (match.missing_skills if match else None) or []
    jd_skills = matched_skills + missing_skills

    resume_text_parts = [
        resume_data.get("summary", ""),
        " ".join(resume_data.get("skills", [])),
        requirement.role,
    ]
    for exp in resume_data.get("experience", []):
        resume_text_parts.append(exp.get("role", ""))
        resume_text_parts.append(" ".join(exp.get("bullets", [])))
    resume_text = " ".join(resume_text_parts)

    if jd_skills:
        ats_total, kw_score, role_score, fmt_score, ats_matched, ats_missing = _ats_score(
            jd_skills, resume_text, requirement.role
        )
        generated.ats_score = ats_total
        generated.ats_keyword_score = kw_score
        generated.ats_role_score = role_score
        generated.ats_format_score = fmt_score
        generated.ats_matched_keywords = ats_matched
        generated.ats_missing_keywords = ats_missing

    # Rebuild the actual files from the edited content — same file layout
    # and naming as initial generation, so download links keep working.
    full_name = (consultant.full_name or "").strip()
    name_parts = full_name.split() if full_name else ["", ""]
    first_name = name_parts[0] if name_parts else ""
    last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""

    base_filename_with_ext = _build_resume_filename(
        first_name=first_name,
        last_name=last_name,
        role=requirement.role,
        client_name=requirement.client,
        vendor_name=requirement.vendor,
        years_exp=consultant.total_experience_years,
    )
    base_stem = base_filename_with_ext.removesuffix(".pdf")

    resume_dir = RESUME_UPLOAD_DIR / "generated" / str(consultant.id) / str(requirement.id)
    resume_dir.mkdir(parents=True, exist_ok=True)
    docx_path = resume_dir / f"{base_stem}.docx"
    pdf_path = resume_dir / f"{base_stem}.pdf"

    try:
        _generate_docx(resume_data, docx_path)
        pdf_ok = _convert_to_pdf(docx_path, pdf_path)
    except Exception as exc:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Failed to regenerate DOCX: {exc}")

    s3_pdf_key = None
    if pdf_ok:
        from s3_service import upload_file_to_s3
        s3_pdf_key = f"generated/{consultant.id}/{requirement.id}/{base_stem}.pdf"
        with open(pdf_path, "rb") as f:
            if not upload_file_to_s3(f, s3_pdf_key, "application/pdf"):
                s3_pdf_key = None

    generated.filename = base_filename_with_ext
    generated.docx_path = str(docx_path)
    generated.pdf_path = s3_pdf_key if s3_pdf_key else (str(pdf_path) if pdf_ok else None)
    generated.generation_status = "COMPLETED" if pdf_ok else generated.generation_status

    await db.commit()
    await db.refresh(generated)

    return _build_resume_data_dto(generated, requirement)


@router.get(
    "/api/consultant/requirements/{requirement_id}/resume/history",
    summary="Get all generation attempts — returns ResumeHistoryDTO",
)
async def get_resume_history(
    requirement_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Returns ResumeHistoryDTO matching frontend types/index.ts:
    { requirementId, attempts: [{ attemptNumber, atsScore, generatedAt }] }
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    result = await db.execute(
        select(GeneratedResume)
        .where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement_id,
        )
        .order_by(GeneratedResume.created_at.asc())
    )
    resumes = result.scalars().all()

    attempts = [
        {
            "attemptNumber": r.generation_attempt,
            "atsScore": float(r.ats_score or 0),
            "generatedAt": r.created_at.isoformat() if r.created_at else "",
        }
        for r in resumes
    ]

    return {
        "requirementId": str(requirement_id),
        "attempts": attempts,
    }


@router.get(
    "/api/consultant/requirements/{requirement_id}/resume/download/{file_type}",
    summary="Download resume file (pdf or docx) — serves actual file",
    response_class=FileResponse,
)
async def download_resume(
    requirement_id: int,
    file_type: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Serves the actual DOCX or PDF file.
    Frontend DownloadButtonGroup and ResumeActionCell both link to these URLs.
    file_type must be 'pdf' or 'docx'.
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    if file_type not in ("pdf", "docx"):
        raise HTTPException(status_code=422, detail="file_type must be 'pdf' or 'docx'")

    result = await db.execute(
        select(GeneratedResume).where(
            GeneratedResume.consultant_id == consultant.id,
            GeneratedResume.requirement_id == requirement_id,
            GeneratedResume.is_final == True,
        )
    )
    generated = result.scalars().first()
    if not generated:
        raise HTTPException(status_code=404, detail="No generated resume found.")

    file_path = generated.pdf_path if file_type == "pdf" else generated.docx_path

    # BUG FIX: generated.filename is always stored ending in ".pdf"
    # (see _build_resume_filename — it hardcodes the suffix), and that one
    # stored name gets reused for BOTH the pdf_path and docx_path files.
    # Downloading "pdf" happened to look right by coincidence, but
    # downloading "docx" served real DOCX bytes under a filename ending in
    # ".pdf" — both downloads then land in the Downloads folder with the
    # same base name, making it easy to open the wrong one and see Word
    # content where a PDF was expected. Force the extension to match what
    # was actually requested, regardless of what's stored.
    stored_name = generated.filename or Path(file_path).name
    stem = stored_name.rsplit(".", 1)[0] if "." in stored_name else stored_name
    correct_filename = f"{stem}.{file_type}"

    if file_path and not Path(file_path).exists():
        from s3_service import download_file_from_s3
        body, content_type = download_file_from_s3(file_path)
        if body:
            from fastapi.responses import Response
            return Response(
                content=body,
                media_type=content_type or "application/pdf",
                headers={
                    "Content-Disposition":
                        f'attachment; filename="{correct_filename}"'
                },
            )

    if not file_path or not Path(file_path).exists():
        raise HTTPException(
            status_code=404,
            detail=f"{file_type.upper()} file not available. "
                   + ("Install LibreOffice on the server for PDF generation." if file_type == "pdf" else ""),
        )

    media_type = (
        "application/pdf" if file_type == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=correct_filename,
    )


@router.get(
    "/api/recruiter/consultants/{consultant_id}/requirements/{requirement_id}/resume",
    summary="Get resume data for a consultant (recruiter/admin view) — returns ResumeDataDTO",
)
async def get_resume_recruiter_view(
    consultant_id: int,
    requirement_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Recruiter/Admin view. Enforces recruiter→consultant mapping for RECRUITER role.
    Returns same ResumeDataDTO shape as the consultant endpoint.
    """
    _require_role(current_user, "RECRUITER", "ADMIN")

    if current_user.role == "RECRUITER":
        mapping_result = await db.execute(
            select(RecruiterConsultant).where(
                RecruiterConsultant.recruiter_id == current_user.id,
                RecruiterConsultant.consultant_id == consultant_id,
                RecruiterConsultant.is_active == True,
            )
        )
        if not mapping_result.scalars().first():
            raise HTTPException(status_code=403, detail="Consultant not assigned to this recruiter")

    requirement = await _get_requirement_or_404(db, requirement_id)

    result = await db.execute(
        select(GeneratedResume).where(
            GeneratedResume.consultant_id == consultant_id,
            GeneratedResume.requirement_id == requirement_id,
            GeneratedResume.is_final == True,
        )
    )
    generated = result.scalars().first()
    if not generated:
        raise HTTPException(status_code=404, detail="No generated resume found.")

    return _build_resume_data_dto(generated, requirement)