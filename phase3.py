# phase3.py
# ---------------------------------------------------------------------------
# Phase 3 — Consultant Profiles, Resume Upload, Experience, Recruiter Mapping
#
# Architecture: single flat file in project root, following the same pattern
# as main.py. Reuses get_db, get_current_user, decode_access_token,
# PK_TYPE/FK_TYPE, ArrayTextColumn, JSONBColumn from existing modules.
#
# New endpoints:
#
#   Consultant CRUD
#   POST   /api/consultant/profile              create consultant profile
#   GET    /api/consultant/profile              get own profile (consultant)
#   PUT    /api/consultant/profile              update own profile (consultant)
#   GET    /api/consultants/{id}                get any consultant (admin/recruiter)
#   PUT    /api/consultants/{id}                update any consultant (admin/recruiter)
#   GET    /api/consultants                     list consultants (paginated)
#   POST   /api/admin/consultants               create consultant (admin)
#   PATCH  /api/admin/consultants/{id}/deactivate  soft-delete
#   PATCH  /api/admin/consultants/{id}/activate    reactivate
#
#   Resume
#   POST   /api/consultant/resume/upload        upload base resume (DOCX/PDF)
#   GET    /api/consultant/resume               get resume metadata
#   DELETE /api/consultant/resume               remove resume (admin)
#
#   Experience
#   GET    /api/consultant/experience                     list entries
#   POST   /api/consultant/experience                     add entry
#   PUT    /api/consultant/experience/{id}                full update
#   DELETE /api/consultant/experience/{id}                delete
#   PATCH  /api/consultant/experience/reorder             save sort order
#
#   Recruiter ↔ Consultant Mapping
#   GET    /api/recruiter/consultants                     my consultants
#   POST   /api/recruiter/consultants                     assign consultant
#   DELETE /api/recruiter/consultants/{id}                unassign
#   GET    /api/admin/consultants/{id}/recruiters         list recruiters (admin)
#   PUT    /api/admin/consultants/{id}/recruiters         set recruiters (admin)
# ---------------------------------------------------------------------------

from __future__ import annotations

import asyncio
import io
import logging
import math
import os
import re
import secrets
import string
import uuid
from datetime import datetime, date
from pathlib import Path
from typing import List, Optional, Dict, Any

from fastapi import APIRouter, Depends, File, HTTPException, Query, Request, UploadFile, status, BackgroundTasks
from pydantic import BaseModel, EmailStr, Field, field_validator, model_validator
from sqlalchemy import func, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select

from database import get_db
from models import (
    User,
    Consultant,
    RecruiterConsultant,
    ConsultantExperience,
    Application,
    Resume,
    PK_TYPE,
    FK_TYPE,
)

# ---------------------------------------------------------------------------
# Re-use get_current_user and decode_access_token from main.py
# Imported here at module level — no circular dependency because phase3.py
# does NOT import the FastAPI app instance, only the utility functions.
# ---------------------------------------------------------------------------
from auth import get_current_user, decode_access_token
from s3_service import upload_file_to_s3, delete_file_from_s3

logger = logging.getLogger(__name__)

router = APIRouter()

# Per-consultant background-rematch coalescing guard — see
# _rematch_in_background in update_own_profile below. Per-process only;
# would need a DB/Redis-backed lock for a multi-worker deployment.
_rematch_in_progress: set[int] = set()
_rematch_dirty: set[int] = set()

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# FEATURE CHANGE: only .docx is accepted now, no PDF — matches the same
# restriction applied to the "My Resumes" upload flow (resume_router.py).
ALLOWED_RESUME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
MAX_RESUME_BYTES = 10 * 1024 * 1024  # 10 MB
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads/resumes"))

# ---------------------------------------------------------------------------
# Pydantic schemas — matching frontend DTO contracts exactly
# (ConsultantProfileDTO, ExperienceDTO from frontend types/index.ts)
# ---------------------------------------------------------------------------

# ── Profile ─────────────────────────────────────────────────────────────────

class EducationEntryRequest(BaseModel):
    # BUG FIX: every field here was Optional with no validation, and
    # ProfileUpdateRequest.education defaulted to [] with no min_length —
    # matching the frontend's schema gap (see ProfileFormSchema.education
    # in schemas/index.ts), so an empty education list, or an entry with
    # blank degree/institution/year, always saved successfully with no
    # server-side check at all, independent of whatever the UI did or
    # didn't enforce.
    degree: str = Field(..., min_length=1)
    institution: str = Field(..., min_length=1)
    year: str = Field(..., min_length=1)
    details: Optional[str] = None


class EducationEntryResponse(BaseModel):
    """BUG FIX (GET /api/consultants 500 — 'education.1.institution:
    String should have at least 1 character'): ProfileResponse.education
    used to be typed List[EducationEntryRequest] — reusing the STRICT
    request-validation class (min_length=1 on every field, meant to
    reject an incomplete save from the consultant's own profile form) as
    the response schema too. Real stored data doesn't carry that
    guarantee — e.g. an education row added through the admin editor
    with institution left blank — so building the response for any
    consultant with one such row threw a validation error and 500'd the
    ENTIRE list endpoint, not just that consultant. Response serialization
    should describe what the data actually is, not what a future write
    must satisfy; those are different concerns even though the shape
    looks identical. Every field optional here, matching
    phase_users_schema.py's EducationEntryDTO (the admin-side equivalent
    response schema, which never had this problem for the same reason).
    """
    degree: Optional[str] = None
    institution: Optional[str] = None
    year: Optional[str] = None
    details: Optional[str] = None


class ProfileUpdateRequest(BaseModel):
    fullName: str = Field(..., min_length=1, max_length=200)
    # BUG FIX ("saving profile with empty Location/Phone/LinkedIn URL/
    # Preferred Roles/Preferred Locations/Total Experience"): every one of
    # these fields shows the required red asterisk in ProfileForm.tsx and
    # already renders its own errors.<field> message correctly, but each
    # was Optional here with no server-side validation at all — matching
    # the same schema gap fixed on the frontend side (see
    # ProfileFormSchema in schemas/index.ts) for phone, linkedInUrl,
    # preferredRoles, preferredLocations, and totalExperienceYears, all of
    # which had a `.or(z.literal(""))` (or equivalent) escape hatch there.
    location: str = Field(..., min_length=2, max_length=100)
    phone: str = Field(..., pattern=r"^\+?[\d\s\-().]{7,20}$")
    # (education itself is declared further down in this class — see
    # the min_length=1 added there, paired with EducationEntryRequest's
    # own per-field requirements above.)
    linkedInUrl: str = Field(..., min_length=1)
    # BUG FIX ("saving profile with zero Primary/Secondary Skills, or with
    # no Work Authorisation selected"): both skill sections show a
    # required asterisk in SkillTagInput.tsx ("Primary Skills *" /
    # "Secondary Skills *"), and Work Authorisation shows one in
    # WorkAuthSelect.tsx — but neither component uses react-hook-form or
    # any Zod schema at all (they call updateMutation.mutate(...) with a
    # raw object directly), so the SkillsSchema/WorkAuthFormSchema exports
    # in schemas/index.ts were never actually wired to either of them —
    # dead code, not real enforcement. Unlike
    # EmploymentTypeCheckboxGroup.tsx (which already self-guards with
    # `if (updated.length === 0) return;` before ever calling mutate),
    # SkillTagInput.tsx's removeSkill had no equivalent guard, and every
    # OTHER component's auto-save (ProfileForm, WorkAuthSelect,
    # EmploymentTypeCheckboxGroup) passes primarySkills/secondarySkills/
    # workAuth straight through from the current profile prop unchanged
    # — so a brand-new consultant who never touches Skills or Work
    # Authorisation at all could still "successfully" save every other
    # section indefinitely while silently carrying forward empty
    # skills / a null work auth, with the UI's required asterisks never
    # actually being enforced by anything.
    primarySkills: List[str] = Field(..., min_length=1)
    secondarySkills: List[str] = Field(..., min_length=1)
    workAuth: str = Field(...)
    # BUG FIX: default=["C2C"] was itself an now-invalid legacy value once
    # the allowed set below was narrowed to FULL_TIME/CONTRACT — required
    # here with no default, matching the pattern used for the other
    # fields that were made required earlier (title, education, etc.).
    employmentTypes: List[str] = Field(..., min_length=1)
    preferredRoles: str = Field(..., min_length=1, max_length=200)
    preferredLocations: str = Field(..., min_length=1, max_length=200)
    totalExperienceYears: float = Field(..., ge=0, le=60)
    # BUG FIX: these three were never collectable anywhere — the
    # "Profile incomplete" check (resume_validation.py) has always
    # required them, but there was no form field, no request field, and
    # (for title/summary/education) no storage column at all. Stored in
    # User.resume_info (see update_own_profile below) rather than new
    # Consultant columns, to avoid a migration.
    # BUG FIX ("saving profile with empty Target Title/Role"): title was
    # Optional with no validation, matching the frontend's schema gap
    # (see ProfileFormSchema.title in schemas/index.ts) — an empty string
    # always saved successfully with no server-side check, independent of
    # the UI's required asterisk.
    title: str = Field(..., min_length=2, max_length=150)
    summary: Optional[str] = None
    education: List[EducationEntryRequest] = Field(..., min_length=1)
    resumeRichText: Optional[str] = None
    # BUG FIX ("switching quickly between Work Auth options sometimes
    # ends up on the wrong one"): client-generated, strictly increasing
    # write sequence (Date.now() at click time) — see
    # Consultant.last_profile_write_seq in models.py for why this exists.
    # Optional so any caller that doesn't send it just skips the
    # staleness check (old behavior).
    clientWriteSeq: Optional[int] = None

    # BUG FIX: workAuth is now required (str, not Optional[str], above) —
    # this validator's `v is not None` branch is unreachable now, since
    # Pydantic rejects a missing/None value before this ever runs, but the
    # membership check itself still matters (a required field isn't
    # automatically a valid enum member).
    @field_validator("workAuth")
    @classmethod
    def validate_work_auth(cls, v):
        # BUG FIX: these exact strings (once normalized) must line up with
        # phase4.py's WORK_AUTH_BATCH_1/2/3 sets — get_batch() only strips
        # spaces and hyphens before comparing, so an underscore-separated
        # value here would silently fail to match any batch and get every
        # consultant wrongly rejected at Stage 2 of validate_match().
        valid = {"F1", "STEM OPT", "H1B", "USC", "GC", "GC EAD", "L1", "TN", "U Visa"}
        if v not in valid:
            raise ValueError(f"workAuth must be one of {', '.join(sorted(valid))}")
        return v

    # BUG FIX: linkedInUrl is now required (min_length=1 above), but the
    # frontend's additional check that it's actually a LinkedIn link
    # (ProfileFormSchema's .refine(val.includes("linkedin.com"))) had no
    # backend equivalent — any non-empty string, LinkedIn or not, saved
    # successfully via a direct API call.
    @field_validator("linkedInUrl")
    @classmethod
    def validate_linkedin_url(cls, v):
        if not re.match(r"^https?://", v):
            raise ValueError("linkedInUrl must be a valid URL")
        if "linkedin.com" not in v:
            raise ValueError("linkedInUrl must be a LinkedIn URL")
        return v

    # BUG FIX ("selecting Contract and hard-refreshing unselects it", then
    # "every save fails with Invalid employmentTypes: ['C2C']"): narrowing
    # the allowed set to just FULL_TIME/CONTRACT fixed Contract itself,
    # but exposed a second bug — any consultant created via
    # admin_create_consultant (whose own AdminConsultantCreateRequest
    # still allows C2C/W2/1099/FULL_TIME/CONTRACT) already has one of
    # those legacy values stored. Every other component on this page
    # (ProfileForm, SkillTagInput, WorkAuthSelect,
    # EmploymentTypeCheckboxGroup) re-sends the CURRENT employmentTypes
    # array untouched alongside whatever field it's actually changing —
    # so a hard reject here blocked every save on the entire page for
    # that consultant, not just the Employment Type toggle. Filtering out
    # unrecognized legacy values instead lets those other saves succeed
    # normally and self-heals the stale data the next time employmentTypes
    # itself is genuinely touched; only raise if nothing valid survives
    # the filter, since it's still a required field.
    #
    # BUG FIX ("Failed to save — changes rolled back" on ANY profile edit
    # for consultants with only legacy employmentTypes): rejecting here
    # whenever nothing survives the filter looked right in isolation, but
    # this validator has no DB access — it can't tell "the person just
    # cleared Employment Type" from "ProfileForm/WorkAuthSelect/
    # SkillTagInput re-sent the CURRENT value unchanged, and that value
    # happens to be legacy-only (C2C/W2/1099 — still allowed by
    # AdminConsultantCreateRequest, so plenty of existing consultants have
    # exactly this)." The latter is the common case, and raising blocked
    # every save on the page for those consultants, not just Employment
    # Type. Just filter here; whether an empty result should overwrite
    # the stored value is decided in the endpoint, which has the existing
    # row and CAN tell the difference (see _resolve_employment_types).
    @field_validator("employmentTypes")
    @classmethod
    def validate_employment_types(cls, v):
        allowed = {"FULL_TIME", "CONTRACT"}
        return list(dict.fromkeys(t for t in v if t in allowed))


# BUG FIX (app crashed on startup): update_consultant_by_id below declares
# its request body as AdminConsultantUpdateRequest, but that class didn't
# exist anywhere in this file. Because `from __future__ import annotations`
# is active at the top of this module, that missing name didn't raise
# immediately at function-definition time — but FastAPI resolves every
# route's type hints (via typing.get_type_hints) to build its request-body
# validator the moment this router is registered, which happens as soon as
# main.py imports it at startup. That resolution raised
# `NameError: name 'AdminConsultantUpdateRequest' is not defined` and took
# the entire app down before it could even start serving requests — not
# just this one endpoint.
#
# Defined here to mirror ProfileUpdateRequest's field names and validators
# for every field update_consultant_by_id actually assigns below, but
# omitting title/summary/clientWriteSeq — this is an admin or recruiter
# editing someone ELSE's profile, not the consultant's own profile form,
# and the endpoint never reads those three fields.
class AdminConsultantUpdateRequest(BaseModel):
    fullName: str = Field(..., min_length=1, max_length=200)
    location: str = Field(..., min_length=2, max_length=100)
    phone: str = Field(..., pattern=r"^\+?[\d\s\-().]{7,20}$")
    linkedInUrl: str = Field(..., min_length=1)
    primarySkills: List[str] = Field(..., min_length=1)
    secondarySkills: List[str] = Field(..., min_length=1)
    workAuth: str = Field(...)
    employmentTypes: List[str] = Field(..., min_length=1)
    preferredRoles: str = Field(..., min_length=1, max_length=200)
    preferredLocations: str = Field(..., min_length=1, max_length=200)
    totalExperienceYears: float = Field(..., ge=0, le=60)
    education: List[EducationEntryRequest] = Field(..., min_length=1)
    resumeRichText: Optional[str] = None

    @field_validator("workAuth")
    @classmethod
    def validate_work_auth(cls, v):
        valid = {"F1", "STEM OPT", "H1B", "USC", "GC", "GC EAD", "L1", "TN", "U Visa"}
        if v not in valid:
            raise ValueError(f"workAuth must be one of {', '.join(sorted(valid))}")
        return v

    @field_validator("linkedInUrl")
    @classmethod
    def validate_linkedin_url(cls, v):
        if not re.match(r"^https?://", v):
            raise ValueError("linkedInUrl must be a valid URL")
        if "linkedin.com" not in v:
            raise ValueError("linkedInUrl must be a LinkedIn URL")
        return v

    @field_validator("employmentTypes")
    @classmethod
    def validate_employment_types(cls, v):
        allowed = {"FULL_TIME", "CONTRACT"}
        return list(dict.fromkeys(t for t in v if t in allowed))


class ProfileResponse(BaseModel):
    model_config = {"from_attributes": True}
    id: str
    fullName: Optional[str] = None
    email: Optional[str] = None
    location: Optional[str] = None
    phone: Optional[str] = None
    linkedInUrl: Optional[str] = None
    primarySkills: List[str] = []
    secondarySkills: List[str] = []
    workAuth: Optional[str] = None
    employmentTypes: List[str] = []
    resume: Optional[dict] = None
    experienceCount: int = 0
    gmailConnected: bool = False
    baseResumeUploaded: bool = False
    atsScore: float = 0.0
    status: str = "ACTIVE"
    preferredRoles: Optional[str] = None
    preferredLocations: Optional[str] = None
    totalExperienceYears: Optional[float] = None
    resume_info: Optional[dict] = None
    availabilityStatus: Optional[str] = None
    createdAt: Optional[str] = None
    # BUG FIX: previously missing — recruiter roster grid (ConsultantCard)
    # and the recruiter/admin detail views hardcoded these to 0 on the
    # frontend because this endpoint never returned them. Now computed for
    # real in _consultant_to_profile_response below.
    profileCompleteness: int = 0
    activeApplicationsCount: int = 0
    # BUG FIX: read back from User.resume_info — see ProfileUpdateRequest.
    title: Optional[str] = None
    summary: Optional[str] = None
    education: List[EducationEntryResponse] = []
    resumeRichText: Optional[str] = None


# ---------------------------------------------------------------------------
# Admin "Add Consultant" request/response — field names, enum values, and
# response shape verified against the actual frontend form component
# (AddConsultantDrawer.tsx): its useCreateConsultant() hook posts snake_case
# keys straight from `form` state, validates work_auth against WORK_AUTHS =
# ["USC","GC","H1B","OPT","CPT","EAD","TN","Other"] and employment_prefs
# against EMP_PREFS = ["C2C","W2","1099","FULL_TIME","CONTRACT"], and reads
# result.message / result.temp_password back from CreateConsultantResponseDTO
# to show the "temporary password — shown once" panel. This DTO intentionally
# does NOT follow the camelCase convention ProfileUpdateRequest uses — its
# only consumer is this one admin form, so it matches that form's payload
# shape exactly, same as the per-endpoint casing convention phase5.py uses.
# ---------------------------------------------------------------------------

_ADMIN_WORK_AUTHS = {"USC", "GC", "H1B", "OPT", "CPT", "EAD", "TN", "Other"}
_ADMIN_EMPLOYMENT_PREFS = {"C2C", "W2", "1099", "FULL_TIME", "CONTRACT"}


class AdminConsultantCreateRequest(BaseModel):
    name: str = Field(..., min_length=2, max_length=100)
    email: EmailStr
    work_auth: str
    employment_prefs: List[str] = Field(..., min_length=1)
    primary_skills: str = ""
    recruiter_id: Optional[str] = None
    phone: Optional[str] = Field(None, max_length=30)
    current_location: Optional[str] = None
    preferred_locations: Optional[str] = None
    availability_status: Optional[str] = None
    total_experience_years: Optional[float] = Field(None, ge=0, le=60)
    secondary_skills: Optional[str] = None
    preferred_roles: Optional[str] = None
    resume_info: Optional[dict] = None
    resume_rich_text: Optional[str] = None
    linkedin_url: Optional[str] = None
    education: List[EducationEntryRequest] = []

    @field_validator("email")
    @classmethod
    def normalise_email(cls, v):
        return v.lower().strip()

    @field_validator("work_auth")
    @classmethod
    def validate_work_auth(cls, v):
        if v not in _ADMIN_WORK_AUTHS:
            raise ValueError(f"work_auth must be one of {sorted(_ADMIN_WORK_AUTHS)}")
        return v

    @field_validator("employment_prefs")
    @classmethod
    def validate_employment_prefs(cls, v):
        invalid = [t for t in v if t not in _ADMIN_EMPLOYMENT_PREFS]
        if invalid:
            raise ValueError(f"Invalid employment_prefs: {invalid}")
        return list(dict.fromkeys(v))


class CreateConsultantResponse(BaseModel):
    """Matches frontend CreateConsultantResponseDTO — the drawer reads
    result.message and result.temp_password directly."""
    message: str
    temp_password: str
    consultant_id: str
    name: str
    email: str


class ConsultantListResponse(BaseModel):
    data: List[ProfileResponse]
    total: int
    page: int
    page_size: int
    total_pages: int


class ResumeUploadResponse(BaseModel):
    resume: dict   # matches ResumeInfoDTO: { filename, uploadedAt, sizeBytes }


# ── Experience ───────────────────────────────────────────────────────────────

class ExperienceMonthYear(BaseModel):
    month: int = Field(..., ge=1, le=12)
    year: int = Field(..., ge=1970, le=2100)


class ExperienceRequest(BaseModel):
    clientName: str = Field(..., min_length=1, max_length=200)
    implementationPartner: Optional[str] = Field(None, max_length=200)
    roleTitle: str = Field(..., min_length=1, max_length=200)
    startDate: ExperienceMonthYear
    endDate: Optional[ExperienceMonthYear] = None
    isPresent: bool = False
    location: Optional[str] = None
    workMode: Optional[str] = Field(None, pattern="^(REMOTE|ONSITE|HYBRID|TRAVEL_REQUIRED)$")
    workModeDetail: Optional[str] = None
    technologies: List[str] = []
    responsibilities: Optional[str] = None
    achievements: Optional[str] = None
    sortOrder: int = 0

    @model_validator(mode="after")
    def validate_dates(self):
        if self.isPresent:
            self.endDate = None
        return self


class ExperienceResponse(BaseModel):
    model_config = {"from_attributes": True}
    id: str
    clientName: str
    implementationPartner: Optional[str] = None
    roleTitle: str
    startDate: Optional[ExperienceMonthYear] = None
    endDate: Optional[ExperienceMonthYear] = None
    isPresent: bool = False
    location: Optional[str] = None
    workMode: Optional[str] = None
    workModeDetail: Optional[str] = None
    technologies: List[str] = []
    responsibilities: Optional[str] = None
    achievements: Optional[str] = None
    sortOrder: int = 0


class ReorderRequest(BaseModel):
    orderedIds: List[str] = Field(..., min_length=1)


# ── Recruiter Mapping ────────────────────────────────────────────────────────

class AssignConsultantRequest(BaseModel):
    consultantId: int = Field(..., gt=0)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _require_role(user: User, *roles: str) -> None:
    if user.role not in roles:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail=f"Requires role: {list(roles)}",
        )


async def _get_consultant_or_404(db: AsyncSession, consultant_id: int) -> Consultant:
    result = await db.execute(select(Consultant).where(Consultant.id == consultant_id))
    c = result.scalars().first()
    if not c:
        raise HTTPException(status_code=404, detail="Consultant not found")
    return c


async def _assert_recruiter_mapped(db: AsyncSession, recruiter_id: int, consultant_id: int) -> None:
    result = await db.execute(
        select(RecruiterConsultant).where(
            RecruiterConsultant.recruiter_id == recruiter_id,
            RecruiterConsultant.consultant_id == consultant_id,
            RecruiterConsultant.is_active == True,
        )
    )
    if not result.scalars().first():
        raise HTTPException(status_code=403, detail="Consultant not assigned to this recruiter")


async def _get_consultant_for_user(db: AsyncSession, user: User) -> Consultant:
    """Return the Consultant row linked to a CONSULTANT-role user."""
    result = await db.execute(select(Consultant).where(Consultant.user_id == user.id))
    c = result.scalars().first()
    if not c:
        raise HTTPException(status_code=404, detail="Consultant profile not found for this user")
    return c


async def _consultant_to_profile_response(
    db: AsyncSession, c: Consultant, experience_count: int = 0, *, include_resume_size: bool = True
) -> ProfileResponse:
    # BUG FIX: title/summary/education/linkedin have no Consultant
    # column — read them from the linked User.resume_info JSON instead,
    # the same blob update_own_profile now writes them into.
    resume_info: Dict[str, Any] = {}
    if c.user_id:
        user_result = await db.execute(select(User.resume_info).where(User.id == c.user_id))
        resume_info = user_result.scalar_one_or_none() or {}
    """Map ORM Consultant → ProfileResponse matching frontend ConsultantProfileDTO."""
    primary = [s.strip() for s in (c.primary_skills or "").split(",") if s.strip()]
    secondary = [s.strip() for s in (c.secondary_skills or "").split(",") if s.strip()]
    emp_types = c.preferred_employment_types or []

    # PERF FIX ("taking long time to save"): base_resume_file_path is set
    # for virtually every consultant — the background auto-regeneration in
    # resume_router.py's sync_base_resume_text writes it on every profile
    # save, not just a manual override upload. get_s3_file_metadata() calls
    # boto3's SYNCHRONOUS head_object() with no await/thread offload —
    # inside this async function that blocks the entire event loop for the
    # duration of a real network round-trip to Spaces, on every profile
    # fetch, and it was ALSO running as part of the save request itself
    # (update_own_profile/update_consultant_by_id both build their response
    # through this same function). Two fixes:
    #   1. include_resume_size=False on the save endpoints below skips the
    #      S3 call entirely — the person just saved, a moment-stale resume
    #      size is fine, and this removes a whole network round-trip from
    #      the save path (the biggest win for perceived save speed).
    #   2. Everywhere else that still needs it (GET profile, list_consultants)
    #      now runs it via asyncio.to_thread so it no longer blocks every
    #      OTHER concurrent request while it waits on Spaces.
    resume = None
    if c.base_resume_file_path:
        fname = Path(c.base_resume_file_path).name
        size_bytes = 0
        if include_resume_size:
            from s3_service import get_s3_file_metadata
            size_bytes, _content_type = await asyncio.to_thread(get_s3_file_metadata, c.base_resume_file_path)
            size_bytes = size_bytes or 0
        resume = {
            "filename": fname,
            "uploadedAt": c.updated_at.isoformat() if c.updated_at else datetime.utcnow().isoformat(),
            "sizeBytes": size_bytes,
        }

    # BUG FIX: was `float(c.ats_score or 0)` — Consultant.ats_score is a
    # column that is never written anywhere in the codebase (only ever read
    # here and in phase_users_service.py), so it was permanently stuck at
    # its default of 0 and every "ATS Score" in the UI showed 0. The real,
    # meaningful ATS score lives per-generated-resume (Resume.ats_score,
    # written by phase6.py/resume_router.py whenever a tailored resume is
    # built). Use the consultant's most recently generated resume's score
    # instead — same signal the resume screens already show.
    latest_resume_result = await db.execute(
        select(Resume.ats_score)
        .where(Resume.user_id == c.user_id, Resume.ats_score.isnot(None))
        .order_by(Resume.created_at.desc())
        .limit(1)
    )
    latest_ats_score = latest_resume_result.scalar_one_or_none() if c.user_id else None

    # BUG FIX: profileCompleteness and activeApplicationsCount were not
    # returned by this endpoint at all — ConsultantCard (recruiter roster
    # grid) and the recruiter/admin detail pages hardcoded both to 0 on the
    # frontend with a comment noting the backend gap. Completeness formula
    # mirrors ProfileCompletenessBar.tsx / phase_users_service.py so every
    # screen agrees on the same number. Active apps = applications actually
    # SENT for this consultant (same definition phase_users_service.py uses
    # for total_applications_sent).
     
    completeness = 0
    if (c.primary_skills or "").strip() or (c.secondary_skills or "").strip():
        completeness += 30  # Skills
    if experience_count > 0:
        completeness += 25  # Experience
    if c.preferred_employment_types:
        completeness += 20  # Employment type
    if (c.work_authorization or "").strip():
        completeness += 15  # Work auth
    if len((c.current_location or "").strip()) >= 2:
        completeness += 10  # Location

    active_apps_result = await db.execute(
        select(func.count()).select_from(Application).where(
            Application.consultant_id == c.id, Application.status == "SENT"
        )
    )
    active_applications_count = active_apps_result.scalar_one() or 0

    return ProfileResponse(
        id=str(c.id),
        fullName=c.full_name,
        email=c.email,
        location=c.current_location,
        phone=c.phone,
        # Prefer the real Consultant.linkedin_url column (what admin edits
        # write to) — fall back to the legacy resume_info blob only for
        # rows saved before this endpoint wrote the column directly.
        # BUG FIX: `c.linkedin_url or resume_info.get(...)` treated an
        # explicitly-cleared LinkedIn URL (admin/consultant blanked the
        # field, leaving "") the same as "never set" — Python's `or` falls
        # through on any falsy value, not just None — so clearing the field
        # would silently resurrect old, stale data from the legacy blob
        # instead of showing empty. Only fall back when the column is
        # genuinely unset (None), which only happens for rows saved before
        # this column existed.
        linkedInUrl=c.linkedin_url if c.linkedin_url is not None else resume_info.get("linkedin"),
        primarySkills=primary,
        secondarySkills=secondary,
        workAuth=c.work_authorization,
        employmentTypes=emp_types,
        resume=resume,
        experienceCount=experience_count,
        gmailConnected=c.gmail_connected,
        baseResumeUploaded=bool(c.base_resume_file_path),
        atsScore=float(latest_ats_score) if latest_ats_score is not None else 0.0,
        status=c.status,
        preferredRoles=c.preferred_roles,
        preferredLocations=c.preferred_locations,
        resume_info=resume_info,
        title=resume_info.get("title"),
        summary=resume_info.get("summary"),
        # Prefer the real Consultant.education column (what admin now
        # edits) — fall back to the legacy resume_info blob only for rows
        # saved before update_own_profile started writing the column too.
        education=c.education if c.education is not None else resume_info.get("education", []),
        resumeRichText=c.resume_rich_text,
        totalExperienceYears=float(c.total_experience_years) if c.total_experience_years is not None else None,
        availabilityStatus=c.availability_status,
        createdAt=c.created_at.isoformat() if c.created_at else None,
        profileCompleteness=completeness,
        activeApplicationsCount=active_applications_count,
    )


def _exp_to_response(e: ConsultantExperience) -> ExperienceResponse:
    """Map ORM ConsultantExperience → ExperienceResponse matching frontend ExperienceDTO."""
    start_date = None
    if e.start_date:
        start_date = ExperienceMonthYear(month=e.start_date.month, year=e.start_date.year)

    end_date = None
    if not e.is_present and e.end_date:
        end_date = ExperienceMonthYear(month=e.end_date.month, year=e.end_date.year)

    return ExperienceResponse(
        id=str(e.id),
        clientName=e.client_name,
        implementationPartner=e.implementation_partner,
        roleTitle=e.role_title,
        startDate=start_date,
        endDate=end_date,
        isPresent=e.is_present,
        location=e.location,
        workMode=e.work_mode,
        workModeDetail=e.work_mode_detail,
        technologies=e.technologies or [],
        responsibilities=e.responsibilities,
        achievements=e.achievements,
        sortOrder=e.sort_order,
    )

def _save_resume_file(file_bytes: bytes, consultant_id: int, original_filename: str, content_type: str) -> str:
    """Upload file to S3, return the S3 key."""
    import io
    from s3_service import upload_file_to_s3
    ext_map = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/pdf": ".pdf",
    }
    ext = ext_map.get(content_type, ".bin")
    s3_key = f"uploads/resumes/{consultant_id}/{uuid.uuid4().hex}{ext}"
    upload_file_to_s3(io.BytesIO(file_bytes), s3_key, content_type)
    return s3_key


def _delete_file_if_exists(path: str) -> None:
    p = Path(path)
    if p.exists() and p.is_file():
        p.unlink()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts)
        return text.replace("\x00", "").replace("\u0000", "")
    except Exception as exc:
        logger.warning("DOCX text extraction failed: %s", exc)
        return ""


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text.strip())
        extracted = "\n".join(parts)
        return extracted.replace("\x00", "").replace("\u0000", "")
    except Exception as exc:
        logger.warning("PDF text extraction failed: %s", exc)
        return ""


def _extract_resume_text(file_bytes: bytes, content_type: str) -> str:
    text = ""
    if "wordprocessingml" in content_type or content_type == "application/docx":
        text = _extract_text_from_docx(file_bytes)
    elif content_type == "application/pdf":
        text = _extract_text_from_pdf(file_bytes)
    return text.replace("\x00", "").replace("\u0000", "")


def _generate_temp_password(length: int = 12) -> str:
    """Generate a random temporary password containing upper, lower, and digit chars."""
    alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
    while True:
        pwd = "".join(secrets.choice(alphabet) for _ in range(length))
        if (
            any(c.islower() for c in pwd)
            and any(c.isupper() for c in pwd)
            and any(c.isdigit() for c in pwd)
        ):
            return pwd


def _hash_password(password: str) -> str:
    """
    Hash a password for storage. Prefers auth.py's own hasher, so the login
    flow (auth.py) verifies it with the exact same algorithm/config it was
    hashed with; falls back to a local passlib bcrypt context only if
    auth.py doesn't expose one.
    """
    try:
        from auth import get_password_hash
        return get_password_hash(password)
    except ImportError:
        pass
    try:
        from passlib.context import CryptContext
        return CryptContext(schemes=["bcrypt"], deprecated="auto").hash(password)
    except ImportError:
        raise RuntimeError(
            "No password hasher available for admin_create_consultant() — "
            "add get_password_hash() to auth.py, or install passlib[bcrypt]."
        )


def _set_password_on_user(user: User, hashed_password: str) -> None:
    """
    Set the hashed password on a new User instance. The exact column name
    isn't referenced anywhere else in this file, so this checks the common
    candidates on the mapped class rather than guessing — update the
    candidate list if your User model uses a different name.
    """
    for candidate in ("hashed_password", "password_hash", "password"):
        if hasattr(User, candidate):
            setattr(user, candidate, hashed_password)
            return
    raise RuntimeError(
        "Could not find a password column on the User model (checked "
        "hashed_password/password_hash/password) — update "
        "_set_password_on_user() in phase3.py with the correct column name."
    )


# Canonical skill library — 100+ skills with aliases for detection
_SKILL_ALIASES: dict[str, list[str]] = {
    "Python": ["python", "python3"],
    "Java": ["java", "core java"],
    "JavaScript": ["javascript", "js", "es6"],
    "TypeScript": ["typescript", "ts"],
    "C#": ["c#", "csharp"],
    "Go": ["golang", "go"],
    "React": ["react", "react.js", "reactjs"],
    "Angular": ["angular", "angularjs"],
    "Vue.js": ["vue", "vue.js", "vuejs"],
    "Next.js": ["next.js", "nextjs"],
    "Node.js": ["node.js", "nodejs"],
    "FastAPI": ["fastapi"],
    "Django": ["django"],
    "Flask": ["flask"],
    "Spring Boot": ["spring boot", "springboot"],
    "PostgreSQL": ["postgresql", "postgres"],
    "MySQL": ["mysql"],
    "MongoDB": ["mongodb", "mongo"],
    "Redis": ["redis"],
    "Elasticsearch": ["elasticsearch"],
    "AWS": ["aws", "amazon web services"],
    "Azure": ["azure", "microsoft azure"],
    "GCP": ["gcp", "google cloud"],
    "Docker": ["docker"],
    "Kubernetes": ["kubernetes", "k8s"],
    "Terraform": ["terraform"],
    "CI/CD": ["ci/cd", "cicd"],
    "REST API": ["rest api", "restful"],
    "GraphQL": ["graphql"],
    "Microservices": ["microservices"],
    "Machine Learning": ["machine learning", "ml"],
    "SQL": ["sql"],
    "Kafka": ["kafka", "apache kafka"],
    "Spark": ["spark", "apache spark", "pyspark"],
    "Airflow": ["airflow", "apache airflow"],
    "Tailwind": ["tailwind", "tailwindcss"],
    "Redux": ["redux"],
    "SAP": ["sap"],
    "Salesforce": ["salesforce", "sfdc"],
    "ServiceNow": ["servicenow"],
    "Linux": ["linux", "ubuntu"],
    "Ansible": ["ansible"],
    "Jenkins": ["jenkins"],
}

_ALIAS_MAP: dict[str, str] = {
    alias.lower(): canonical
    for canonical, aliases in _SKILL_ALIASES.items()
    for alias in aliases
}


def _detect_skills(text: str) -> list[str]:
    """Return canonical skill names found in text, ordered by first appearance."""
    if not text:
        return []
    lower = text.lower()
    found: dict[str, int] = {}
    for alias, canonical in _ALIAS_MAP.items():
        pos = lower.find(alias)
        if pos != -1 and (canonical not in found or pos < found[canonical]):
            found[canonical] = pos
    return [k for k, _ in sorted(found.items(), key=lambda x: x[1])]


# See the BUG FIX note on ProfileUpdateRequest.validate_employment_types
# above — `cleaned` has already been filtered to {FULL_TIME, CONTRACT} by
# then. An empty result there is ambiguous (genuine clear vs. a legacy
# value just being carried forward untouched); resolve that ambiguity
# here, where we actually have the stored row: only a real, non-empty
# selection overwrites it, otherwise leave whatever was already saved
# alone. Shared by both the consultant self-update and the admin/
# recruiter update endpoints below.
def _resolve_employment_types(existing: Optional[List[str]], cleaned: List[str]) -> List[str]:
    return cleaned if cleaned else (existing or [])


# ---------------------------------------------------------------------------
# Consultant Profile endpoints
# ---------------------------------------------------------------------------

@router.get(
    "/api/consultant/profile",
    response_model=ProfileResponse,
    summary="Get own consultant profile",
)
async def get_own_profile(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """CONSULTANT role only — returns their own profile."""
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    count_result = await db.execute(
        select(func.count()).where(ConsultantExperience.consultant_id == consultant.id)
    )
    exp_count = count_result.scalar_one()
    return await _consultant_to_profile_response(db, consultant, exp_count)


@router.put(
    "/api/consultant/profile",
    response_model=ProfileResponse,
    summary="Update own consultant profile",
)
async def update_own_profile(
    payload: ProfileUpdateRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """CONSULTANT role only — updates their own profile."""
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    # BUG FIX ("switching quickly between Work Auth options sometimes
    # ends up on the wrong one"): see Consultant.last_profile_write_seq in
    # models.py. If this request's sequence number is older than the last
    # one this consultant's row actually committed, a later click's
    # request already won — applying this one now would silently revert
    # that newer value just because this slower request happened to
    # arrive last. Drop it and hand back the current (already-newer)
    # state instead of overwriting it.
    if (
        payload.clientWriteSeq is not None
        and consultant.last_profile_write_seq is not None
        and payload.clientWriteSeq <= consultant.last_profile_write_seq
    ):
        count_result = await db.execute(
            select(func.count()).where(ConsultantExperience.consultant_id == consultant.id)
        )
        exp_count = count_result.scalar_one()
        return await _consultant_to_profile_response(db, consultant, exp_count)

    consultant.full_name = payload.fullName
    consultant.current_location = payload.location
    consultant.phone = payload.phone
    consultant.work_authorization = payload.workAuth
    consultant.primary_skills = ", ".join(payload.primarySkills)
    consultant.secondary_skills = ", ".join(payload.secondarySkills)
    # BUG FIX: this used to only stash linkedInUrl inside User.resume_info
    # (see below) — the admin screens (phase_users_service.py, phase3.py's
    # admin update_consultant) read/write the real Consultant.linkedin_url
    # column instead, so a consultant editing their own LinkedIn URL here
    # never showed up on the admin side. Write the real column too so both
    # directions stay in sync.
    consultant.linkedin_url = payload.linkedInUrl
    consultant.preferred_employment_types = _resolve_employment_types(
        consultant.preferred_employment_types, payload.employmentTypes
    )
    consultant.preferred_roles = payload.preferredRoles
    consultant.preferred_locations = payload.preferredLocations
    consultant.total_experience_years = payload.totalExperienceYears
    # Same fix as linkedin_url above: write the real Consultant.education
    # column (what admin now reads/edits) in addition to resume_info below
    # (kept for resume_validation.py's FIELD_CHECKS / resume generation).
    consultant.education = [e.model_dump() for e in payload.education]
    consultant.resume_rich_text = payload.resumeRichText

    # BUG FIX: this consultant self-service endpoint never wrote to
    # User.resume_info at all — the AI generation eligibility check
    # (resume_validation.py) only ever reads from there, so no amount of
    # filling in this form could ever clear the "Profile incomplete"
    # warning. Sync every field it checks for, using the exact keys
    # FIELD_CHECKS expects. linkedInUrl was ALSO previously accepted by
    # this endpoint's request body but never assigned anywhere — genuinely
    # discarded on every save until now.
    existing_info = dict(current_user.resume_info or {})
    existing_info.update({
        "full_name": payload.fullName,
        "email": current_user.email,
        "phone": payload.phone,
        "linkedin": payload.linkedInUrl,
        "location": payload.location,
        "title": payload.title,
        "summary": payload.summary,
        "years_experience": payload.totalExperienceYears,
        "skills": payload.primarySkills + payload.secondarySkills,
        "education": [e.model_dump() for e in payload.education],
    })
    current_user.resume_info = existing_info

    # Record this request as the newest committed write for the
    # out-of-order-write guard above, so any still-in-flight OLDER
    # request that lands after this one gets dropped instead of
    # reverting this value.
    if payload.clientWriteSeq is not None:
        consultant.last_profile_write_seq = payload.clientWriteSeq

    # Regenerate base_resume_text right away — matching_router.py and the
    # completeness check both read it, and neither should have to wait
    # for a manual visit to the Base Resume editor to see this save.
    # PERF FIX ("saving profile takes a long time"): pass existing_info
    # (the resume_info we just built above) straight through instead of
    # letting sync_base_resume_text re-query User.resume_info from the
    # DB — we already have it in memory, so that query was pure wasted
    # round-trip latency on every save.
    from resume_router import sync_base_resume_text
    await sync_base_resume_text(db, consultant, existing_info, background_tasks)

    await db.commit()
    await db.refresh(consultant)

    # Re-run matching for this consultant so their requirement matches reflect
    # the just-saved profile. Runs in the background with its own DB session so
    # it never blocks or fails the profile save.
    #
    # BUG FIX ("saving one field on My Profile takes 8+ seconds, and a
    # second save right after sits stuck pending"): with 44,000+ open
    # requirements in the system, one rematch pass loops over all of
    # them for this consultant — genuinely takes several seconds. Every
    # field-level auto-save (Work Auth, Skills, Employment Type, etc.)
    # used to independently fire its OWN full rematch via
    # asyncio.create_task. Editing more than one field within a few
    # seconds — completely normal usage — spawned multiple overlapping
    # background tasks for the SAME consultant, all trying to upsert
    # into the same RequirementConsultantMatch rows at once, so they
    # serialized against each other via Postgres row locks instead of
    # Python ever noticing they overlapped. Same coalescing shape as the
    # existing base-resume-sync fix and the Gmail sync mutex fix: track
    # whether a rematch is already running for this consultant; if a new
    # save lands mid-run, don't start a second one — just flag that one
    # more pass is needed once the current one finishes, so N rapid
    # saves cost at most 2 rematch passes instead of N concurrent ones.
    # NOTE: this in-memory guard is per-process — fine for a single
    # uvicorn worker, but would need a DB- or Redis-backed lock instead
    # if this ever runs multi-worker.
    consultant_id = consultant.id

    async def _rematch_in_background(cid: int):
        from database import AsyncSessionLocal
        from phase4 import match_consultant
        from matching_router import run_matching_for_consultant
        try:
            while True:
                _rematch_dirty.discard(cid)
                async with AsyncSessionLocal() as bg_session:
                    await match_consultant(bg_session, cid)
                    # COVERAGE GAP FIX (not a matching-condition change):
                    # this only ever refreshed Pipeline A (the admin
                    # Requirements page's match count). Pending
                    # Applications (Pipeline B, the JobMatch table) never
                    # got refreshed when a consultant updated their
                    # profile — even when the update was specifically to
                    # fix a gap keeping them from matching something.
                    # Same session, same trigger, second pipeline.
                    await run_matching_for_consultant(bg_session, cid)
                # If another save landed while this pass was running, do
                # exactly one more pass to pick up its latest data — this
                # pass started with whatever was saved BEFORE it began,
                # so it can't have already covered that save.
                if cid not in _rematch_dirty:
                    break
        except Exception as e:
            logger.error("Background auto-match failed for consultant_id=%s: %s", cid, e)
            from error_logger import log_db_error
            await log_db_error(
                stage="background_auto_match",
                error=e,
                source_type="consultant",
                source_id=str(cid),
            )
        finally:
            _rematch_in_progress.discard(cid)

    if consultant_id in _rematch_in_progress:
        # A rematch is already running for this consultant — just mark
        # that it needs one more pass when it's done, instead of
        # spawning a second task that would fight the first one over
        # the same rows.
        _rematch_dirty.add(consultant_id)
    else:
        _rematch_in_progress.add(consultant_id)
        asyncio.create_task(_rematch_in_background(consultant_id))

    count_result = await db.execute(
        select(func.count()).where(ConsultantExperience.consultant_id == consultant.id)
    )
    exp_count = count_result.scalar_one()
    # PERF FIX: skip the S3 head_object round-trip on the save response
    # itself — see _consultant_to_profile_response's include_resume_size
    # note. A subsequent GET (page refresh, revisit) still returns the
    # real size.
    return await _consultant_to_profile_response(db, consultant, exp_count, include_resume_size=False)


@router.get(
    "/api/consultants",
    response_model=ConsultantListResponse,
    summary="List consultants (admin sees all, recruiter sees assigned)",
)
async def list_consultants(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    status: Optional[str] = Query(None),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN", "RECRUITER")

    query = select(Consultant)

    if current_user.role == "RECRUITER":
        assigned = select(RecruiterConsultant.consultant_id).where(
            RecruiterConsultant.recruiter_id == current_user.id,
            RecruiterConsultant.is_active == True,
        )
        query = query.where(Consultant.id.in_(assigned))

    if status:
        if status not in Consultant.VALID_STATUSES:
            raise HTTPException(
                422,
                f"status must be one of {sorted(Consultant.VALID_STATUSES)}",
            )

        query = query.where(Consultant.status == status)

    else:
        from sqlalchemy import or_

        query = query.where(
            Consultant.status == "ACTIVE"
        ).where(
            or_(
                Consultant.user_id.is_(None),
                Consultant.user_id.in_(
                    select(User.id).where(User.is_authorized == True)
                ),
            )
        )

    total = (
        await db.execute(
            select(func.count()).select_from(query.subquery())
        )
    ).scalar_one()

    rows = (
        await db.execute(
            query
            .order_by(Consultant.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
    ).scalars().all()

    # Batch-count experience rows for this whole page in one query instead
    # of N+1 — profileCompleteness now factors in Experience, so leaving
    # this defaulted to 0 here would show every consultant on this roster
    # as missing Experience regardless of their real data.
    exp_counts: Dict[int, int] = {}
    if rows:
        exp_rows = (await db.execute(
            select(ConsultantExperience.consultant_id, func.count())
            .where(ConsultantExperience.consultant_id.in_([c.id for c in rows]))
            .group_by(ConsultantExperience.consultant_id)
        )).all()
        exp_counts = {cons_id: cnt for cons_id, cnt in exp_rows}

    data = [
        await _consultant_to_profile_response(db, c, exp_counts.get(c.id, 0))
        for c in rows
    ]
    return ConsultantListResponse(

        data=data,
        total=total,
        page=page,
        page_size=page_size,
        total_pages=math.ceil(total / page_size) if page_size else 0,
    )


@router.get(
    "/api/consultants/{consultant_id}",
    response_model=ProfileResponse,
    summary="Get consultant by ID (admin or assigned recruiter)",
)
async def get_consultant_by_id(
    consultant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN", "RECRUITER")
    if current_user.role == "RECRUITER":
        await _assert_recruiter_mapped(db, current_user.id, consultant_id)

    consultant = await _get_consultant_or_404(db, consultant_id)
    count_result = await db.execute(
        select(func.count()).where(ConsultantExperience.consultant_id == consultant_id)
    )
    exp_count = count_result.scalar_one()
    return await _consultant_to_profile_response(db, consultant, exp_count)


@router.put(
    "/api/consultants/{consultant_id}",
    response_model=ProfileResponse,
    summary="Update consultant profile (admin or assigned recruiter)",
)
async def update_consultant_by_id(
    consultant_id: int,
    payload: AdminConsultantUpdateRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN", "RECRUITER")
    if current_user.role == "RECRUITER":
        await _assert_recruiter_mapped(db, current_user.id, consultant_id)

    consultant = await _get_consultant_or_404(db, consultant_id)

    consultant.full_name = payload.fullName
    consultant.current_location = payload.location
    consultant.phone = payload.phone
    consultant.work_authorization = payload.workAuth
    consultant.primary_skills = ", ".join(payload.primarySkills)
    consultant.secondary_skills = ", ".join(payload.secondarySkills)
    consultant.linkedin_url = payload.linkedInUrl
    consultant.preferred_employment_types = _resolve_employment_types(
        consultant.preferred_employment_types, payload.employmentTypes
    )
    consultant.preferred_roles = payload.preferredRoles
    consultant.preferred_locations = payload.preferredLocations
    consultant.total_experience_years = payload.totalExperienceYears
    consultant.education = [e.model_dump() for e in payload.education]
    consultant.resume_rich_text = payload.resumeRichText

    from resume_router import sync_base_resume_text
    await sync_base_resume_text(db, consultant, None, background_tasks)

    await db.commit()
    await db.refresh(consultant)

    count_result = await db.execute(
        select(func.count()).where(ConsultantExperience.consultant_id == consultant_id)
    )
    exp_count = count_result.scalar_one()
    # PERF FIX: same as update_own_profile — skip the S3 round-trip on the
    # save response itself.
    return await _consultant_to_profile_response(db, consultant, exp_count, include_resume_size=False)


@router.post(
    "/api/admin/consultants",
    response_model=CreateConsultantResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Create a new consultant profile + login (admin only)",
)
async def admin_create_consultant(
    payload: AdminConsultantCreateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Creates both the login (User, role=CONSULTANT) and the Consultant
    profile, optionally assigning the consultant to a recruiter. Returns a
    one-time temporary password matching AddConsultantDrawer.tsx's expected
    CreateConsultantResponseDTO shape ({ message, temp_password, ... }).
    """
    _require_role(current_user, "ADMIN")

    existing_user = await db.execute(select(User).where(User.email == payload.email))
    if existing_user.scalars().first():
        raise HTTPException(409, f"A user with email '{payload.email}' already exists")
    existing_consultant = await db.execute(select(Consultant).where(Consultant.email == payload.email))
    if existing_consultant.scalars().first():
        raise HTTPException(409, "Consultant with email '{payload.email}' already exists")

    recruiter_id_int: Optional[int] = None
    if payload.recruiter_id:
        try:
            recruiter_id_int = int(payload.recruiter_id)
        except (TypeError, ValueError):
            raise HTTPException(422, f"Invalid recruiter_id: {payload.recruiter_id}")
        r = await db.execute(
            select(User).where(User.id == recruiter_id_int, User.role == "RECRUITER", User.is_authorized == True)
        )
        if not r.scalars().first():
            raise HTTPException(404, f"Active recruiter with id={recruiter_id_int} not found")

    temp_password = _generate_temp_password()

    user = User(
        email=payload.email,
        full_name=payload.name,
        role="CONSULTANT",
        is_active=True,
        # resume_info lives on User (not Consultant) — see generate_resume()
        # in resume_router.py, which reads target_user.resume_info to build
        # the AI resume draft. Previously this creation flow had no way to
        # set it at all, so every new consultant started with an empty
        # profile until someone separately edited them via Users → Edit.
        resume_info=payload.resume_info,
    )
    _set_password_on_user(user, _hash_password(temp_password))
    db.add(user)
    await db.flush()  # populate user.id before creating the linked Consultant row

    consultant = Consultant(
        user_id=user.id,
        full_name=payload.name,
        email=payload.email,
        phone=payload.phone,
        work_authorization=payload.work_auth,
        preferred_employment_types=payload.employment_prefs,
        primary_skills=payload.primary_skills or "",
        secondary_skills=payload.secondary_skills or "",
        status="ACTIVE",
        preferred_roles=payload.preferred_roles,
        preferred_locations=payload.preferred_locations,
        current_location=payload.current_location,
        total_experience_years=payload.total_experience_years,
        linkedin_url=payload.linkedin_url,
        education=[e.model_dump() for e in payload.education],
        resume_rich_text=payload.resume_rich_text,
    )
    # availability_status isn't referenced elsewhere in this file, so only
    # set it if the model actually defines that column.
    if hasattr(Consultant, "availability_status"):
        consultant.availability_status = payload.availability_status

    db.add(consultant)
    await db.flush()  # populate consultant.id before the recruiter mapping

    if recruiter_id_int is not None:
        db.add(RecruiterConsultant(
            recruiter_id=recruiter_id_int,
            consultant_id=consultant.id,
            is_active=True,
        ))

    await db.commit()
    await db.refresh(consultant)

    logger.info("Admin %s created consultant id=%s email=%s", current_user.email, consultant.id, consultant.email)

    return CreateConsultantResponse(
        message=f"Consultant '{consultant.full_name}' created successfully.",
        temp_password=temp_password,
        consultant_id=str(consultant.id),
        name=consultant.full_name or "",
        email=consultant.email or "",
    )


@router.patch(
    "/api/admin/consultants/{consultant_id}/deactivate",
    response_model=ProfileResponse,
    summary="Deactivate a consultant (admin only)",
)
async def deactivate_consultant(
    consultant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN")
    consultant = await _get_consultant_or_404(db, consultant_id)
    consultant.status = "INACTIVE"
    await db.commit()
    await db.refresh(consultant)
    logger.info("Admin %s deactivated consultant id=%s", current_user.email, consultant_id)
    return await _consultant_to_profile_response(db, consultant)


@router.patch(
    "/api/admin/consultants/{consultant_id}/activate",
    response_model=ProfileResponse,
    summary="Reactivate a consultant (admin only)",
)
async def activate_consultant(
    consultant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN")
    consultant = await _get_consultant_or_404(db, consultant_id)
    consultant.status = "ACTIVE"
    await db.commit()
    await db.refresh(consultant)
    return await _consultant_to_profile_response(db, consultant)


# ---------------------------------------------------------------------------
# Resume endpoints
# ---------------------------------------------------------------------------

@router.post(
    "/api/consultant/resume/upload",
    response_model=ResumeUploadResponse,
    summary="Upload base resume — DOCX only, max 10 MB",
)
async def upload_resume(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    CONSULTANT role uploads their own resume.
    Extracts text and detects skills automatically.
    Replaces any previously stored resume.
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    content_type = file.content_type or ""
    if content_type not in ALLOWED_RESUME_TYPES:
        raise HTTPException(400, f"Only DOCX (Word) files are accepted. Got: '{content_type}'")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(400, "Uploaded file is empty")
    if len(file_bytes) > MAX_RESUME_BYTES:
        raise HTTPException(413, "File exceeds 10 MB limit")

    # SECURITY FIX: the checks above only look at the client-supplied
    # Content-Type header, which is trivially spoofable — any file could
    # be uploaded and stored under a .docx name/type as long as the
    # client claims the right header. A real .docx is a ZIP archive
    # (Office Open XML), so first check the actual ZIP magic bytes, then
    # confirm python-docx can genuinely open it as a Word document.
    # BEHAVIOR NOTE: previously, a corrupt/fake file would silently
    # "succeed" here and only fail later inside _extract_resume_text's
    # best-effort try/except (swallowed, producing blank extracted text
    # with no error shown) — so this also surfaces a real error to the
    # consultant instead of silently storing a broken resume.
    if not file_bytes.startswith(b"PK\x03\x04"):
        raise HTTPException(400, "File does not appear to be a valid DOCX (Word) document.")
    try:
        from docx import Document
        Document(io.BytesIO(file_bytes))
    except Exception:
        raise HTTPException(400, "File does not appear to be a valid DOCX (Word) document.")

    # Delete old file if present
    if consultant.base_resume_file_path:
        _delete_file_if_exists(consultant.base_resume_file_path)

    # Store file
    file_path = _save_resume_file(file_bytes, consultant.id, file.filename or "resume", content_type)

    # Extract text (best-effort — never fail the upload)
    extracted_text = _extract_resume_text(file_bytes, content_type)

    # Detect and merge skills
    detected = _detect_skills(extracted_text)
    if detected:
        existing = [s.strip() for s in (consultant.primary_skills or "").split(",") if s.strip()]
        merged = list(dict.fromkeys(existing + detected))
        consultant.primary_skills = ", ".join(merged)

    consultant.base_resume_file_path = file_path
    consultant.base_resume_text = extracted_text
    # A new file replaces the old one's content entirely — clear the
    # structured JSON (base_resume_content) so the Edit Base Resume page's
    # GET /api/resume/base/content backfill re-parses THIS text instead of
    # silently continuing to show whatever was parsed from the previous
    # upload.
    consultant.base_resume_content = None
    await db.commit()
    await db.refresh(consultant)

    logger.info("Resume uploaded consultant_id=%s file=%s skills_detected=%d", consultant.id, file_path, len(detected))

    return ResumeUploadResponse(
        resume={
            "filename": file.filename or Path(file_path).name,
            "uploadedAt": datetime.utcnow().isoformat(),
            "sizeBytes": len(file_bytes),
        }
    )


@router.get(
    "/api/consultant/resume",
    summary="Get own resume metadata",
)
async def get_own_resume(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    if not consultant.base_resume_file_path:
        raise HTTPException(404, "No resume uploaded")

    from s3_service import get_s3_file_metadata
    # PERF FIX: same blocking-boto3-call issue as _consultant_to_profile_
    # response — thread-offload so this doesn't stall the event loop for
    # every other concurrent request while it waits on Spaces.
    size_bytes, _content_type = await asyncio.to_thread(get_s3_file_metadata, consultant.base_resume_file_path)

    return {
        "filename": Path(consultant.base_resume_file_path).name,
        "uploadedAt": consultant.updated_at.isoformat() if consultant.updated_at else None,
        "hasExtractedText": bool(consultant.base_resume_text),
        "extractedTextLength": len(consultant.base_resume_text or ""),
        "sizeBytes": size_bytes or 0,
    }


@router.delete(
    "/api/admin/consultants/{consultant_id}/resume",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete consultant resume (admin only)",
)
async def admin_delete_resume(
    consultant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN")
    consultant = await _get_consultant_or_404(db, consultant_id)

    if consultant.base_resume_file_path:
        _delete_file_if_exists(consultant.base_resume_file_path)

    consultant.base_resume_file_path = None
    consultant.base_resume_text = None
    consultant.base_resume_content = None
    await db.commit()


# ---------------------------------------------------------------------------
# Experience endpoints
# ---------------------------------------------------------------------------

@router.get(
    "/api/consultant/experience",
    response_model=List[ExperienceResponse],
    summary="List own experience entries ordered by sortOrder",
)
async def list_own_experience(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    result = await db.execute(
        select(ConsultantExperience)
        .where(ConsultantExperience.consultant_id == consultant.id)
        .order_by(ConsultantExperience.sort_order.asc())
    )
    return [_exp_to_response(e) for e in result.scalars().all()]


@router.post(
    "/api/consultant/experience",
    response_model=ExperienceResponse,
    status_code=status.HTTP_201_CREATED,
    summary="Add an experience entry",
)
async def create_experience(
    payload: ExperienceRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    # BUG FIX: new experiences are meant to always land at the TOP
    # (sort_order 0) — but simply setting sort_order=0 collides with
    # whatever entry already holds that position (e.g. right after a
    # manual drag-reorder reassigns sort_order 0..N to match the new
    # visual order). Two rows tied at sort_order=0 then depend on the
    # database's tie-breaking (typically insertion order/id), which
    # favored the OLDER entry over the genuinely new one — showing the
    # new entry second instead of first. Shift every existing entry down
    # by one first, so sort_order=0 is actually free before the new row
    # claims it — no tie possible.
    await db.execute(
        update(ConsultantExperience)
        .where(ConsultantExperience.consultant_id == consultant.id)
        .values(sort_order=ConsultantExperience.sort_order + 1)
    )

    exp = ConsultantExperience(
        consultant_id=consultant.id,
        client_name=payload.clientName,
        implementation_partner=payload.implementationPartner,
        role_title=payload.roleTitle,
        start_date=date(payload.startDate.year, payload.startDate.month, 1),
        end_date=date(payload.endDate.year, payload.endDate.month, 1) if payload.endDate else None,
        is_present=payload.isPresent,
        location=payload.location,
        work_mode=payload.workMode,
        work_mode_detail=payload.workModeDetail,
        technologies=payload.technologies,
        responsibilities=payload.responsibilities,
        achievements=payload.achievements,
        sort_order=0,
    )
    db.add(exp)
    await db.flush()

    from resume_router import sync_base_resume_text
    await sync_base_resume_text(db, consultant, None, background_tasks)

    await db.commit()
    await db.refresh(exp)
    return _exp_to_response(exp)


@router.put(
    "/api/consultant/experience/{experience_id}",
    response_model=ExperienceResponse,
    summary="Full update of an experience entry",
)
async def update_experience(
    experience_id: int,
    payload: ExperienceRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    result = await db.execute(
        select(ConsultantExperience).where(
            ConsultantExperience.id == experience_id,
            ConsultantExperience.consultant_id == consultant.id,
        )
    )
    exp = result.scalars().first()
    if not exp:
        raise HTTPException(404, "Experience entry not found")

    exp.client_name = payload.clientName
    exp.implementation_partner = payload.implementationPartner
    exp.role_title = payload.roleTitle
    exp.start_date = date(payload.startDate.year, payload.startDate.month, 1)
    exp.end_date = date(payload.endDate.year, payload.endDate.month, 1) if payload.endDate else None
    exp.is_present = payload.isPresent
    exp.location = payload.location
    exp.work_mode = payload.workMode
    exp.work_mode_detail = payload.workModeDetail
    exp.technologies = payload.technologies
    exp.responsibilities = payload.responsibilities
    exp.achievements = payload.achievements
    exp.sort_order = payload.sortOrder

    from resume_router import sync_base_resume_text
    await sync_base_resume_text(db, consultant, None, background_tasks)

    await db.commit()
    await db.refresh(exp)
    return _exp_to_response(exp)


@router.delete(
    "/api/consultant/experience/{experience_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Delete an experience entry",
)
async def delete_experience(
    experience_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    result = await db.execute(
        select(ConsultantExperience).where(
            ConsultantExperience.id == experience_id,
            ConsultantExperience.consultant_id == consultant.id,
        )
    )
    exp = result.scalars().first()
    if not exp:
        raise HTTPException(404, "Experience entry not found")

    # BUG FIX ("delete returns 204 but the row is still there"): this
    # was indented one level too deep, inside the if not exp: raise
    # block above — since raise exits immediately, and that branch only
    # runs when exp is None, this line was completely unreachable dead
    # code. The endpoint ran to completion, committed successfully, and
    # returned 204 every time — without ever actually deleting anything.
    await db.delete(exp)

    from resume_router import sync_base_resume_text
    await sync_base_resume_text(db, consultant, None, background_tasks)

    await db.commit()


@router.patch(
    "/api/consultant/experience/reorder",
    summary="Save drag-drop sort order — accepts { orderedIds: [str, ...] }",
)
async def reorder_experience(
    payload: ReorderRequest,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Reorder experience entries. orderedIds is a list of experience IDs
    in the desired display order. Each entry's sort_order is set to its
    index in the list.
    """
    _require_role(current_user, "CONSULTANT")
    consultant = await _get_consultant_for_user(db, current_user)

    for idx, exp_id_str in enumerate(payload.orderedIds):
        try:
            exp_id = int(exp_id_str)
        except (ValueError, TypeError):
            raise HTTPException(422, f"Invalid experience id: {exp_id_str}")

        await db.execute(
            update(ConsultantExperience)
            .where(
                ConsultantExperience.id == exp_id,
                ConsultantExperience.consultant_id == consultant.id,
            )
            .values(sort_order=idx)
        )

    from resume_router import sync_base_resume_text
    await sync_base_resume_text(db, consultant, None, background_tasks)

    await db.commit()
    return {"message": f"Reordered {len(payload.orderedIds)} entries"}


# ---------------------------------------------------------------------------
# Recruiter ↔ Consultant Mapping endpoints
# ---------------------------------------------------------------------------

@router.get(
    "/api/recruiter/consultants",
    summary="Get consultants assigned to the current recruiter — returns ConsultantDTO[]",
)
async def get_my_consultants(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Returns ConsultantDTO shape matching frontend recruiter.types.ts:
    { id, name, email, title }
    """
    _require_role(current_user, "RECRUITER", "ADMIN")

    result = await db.execute(
        select(Consultant)
        .join(RecruiterConsultant, RecruiterConsultant.consultant_id == Consultant.id)
        .where(
            RecruiterConsultant.recruiter_id == current_user.id,
            RecruiterConsultant.is_active == True,
        )
        .order_by(Consultant.full_name.asc())
    )
    consultants = result.scalars().all()
    return [
        {
            "id": str(c.id),
            "name": c.full_name or "",
            "email": c.email or "",
            "title": c.preferred_roles.split(",")[0].strip() if c.preferred_roles else "",
        }
        for c in consultants
    ]


@router.post(
    "/api/recruiter/consultants",
    status_code=status.HTTP_201_CREATED,
    summary="Assign a consultant to the current recruiter",
)
async def assign_consultant(
    payload: AssignConsultantRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "RECRUITER", "ADMIN")

    await _get_consultant_or_404(db, payload.consultantId)

    existing_q = await db.execute(
        select(RecruiterConsultant).where(
            RecruiterConsultant.recruiter_id == current_user.id,
            RecruiterConsultant.consultant_id == payload.consultantId,
        )
    )
    existing = existing_q.scalars().first()

    if existing:
        if existing.is_active:
            raise HTTPException(409, "Consultant already assigned to this recruiter")
        existing.is_active = True
        await db.commit()
        await db.refresh(existing)
        return {"id": str(existing.id), "message": "Assignment reactivated"}

    mapping = RecruiterConsultant(
        recruiter_id=current_user.id,
        consultant_id=payload.consultantId,
        is_active=True,
    )
    db.add(mapping)
    await db.commit()
    await db.refresh(mapping)

    logger.info("Recruiter %s assigned consultant_id=%s", current_user.email, payload.consultantId)
    return {"id": str(mapping.id), "message": "Consultant assigned"}


@router.delete(
    "/api/recruiter/consultants/{consultant_id}",
    status_code=status.HTTP_204_NO_CONTENT,
    summary="Unassign a consultant from the current recruiter",
)
async def unassign_consultant(
    consultant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "RECRUITER", "ADMIN")
    await _assert_recruiter_mapped(db, current_user.id, consultant_id)

    await db.execute(
        update(RecruiterConsultant)
        .where(
            RecruiterConsultant.recruiter_id == current_user.id,
            RecruiterConsultant.consultant_id == consultant_id,
        )
        .values(is_active=False)
    )
    await db.commit()
    logger.info("Recruiter %s unassigned consultant_id=%s", current_user.email, consultant_id)


@router.get(
    "/api/admin/consultants/{consultant_id}/recruiters",
    summary="List recruiters assigned to a consultant (admin only)",
)
async def list_recruiters_for_consultant(
    consultant_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN")
    await _get_consultant_or_404(db, consultant_id)

    result = await db.execute(
        select(RecruiterConsultant, User)
        .join(User, User.id == RecruiterConsultant.recruiter_id)
        .where(RecruiterConsultant.consultant_id == consultant_id)
        .order_by(RecruiterConsultant.is_active.desc(), User.full_name.asc())
    )
    return [
        {
            "mappingId": str(m.id),
            "recruiterId": str(m.recruiter_id),
            "recruiterName": u.full_name,
            "recruiterEmail": u.email,
            "isActive": m.is_active,
            "assignedAt": m.created_at.isoformat() if m.created_at else None,
        }
        for m, u in result.all()
    ]


@router.put(
    "/api/admin/consultants/{consultant_id}/recruiters",
    summary="Replace recruiter assignments for a consultant (admin only)",
)
async def set_recruiters_for_consultant(
    consultant_id: int,
    recruiter_ids: List[int],
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "ADMIN")
    consultant = await _get_consultant_or_404(db, consultant_id)

    for rid in recruiter_ids:
        r = await db.execute(
            select(User).where(User.id == rid, User.role == "RECRUITER", User.is_authorized == True)
        )
        if not r.scalars().first():
            raise HTTPException(404, f"Active recruiter with id={rid} not found")

    # Deactivate all current
    await db.execute(
        update(RecruiterConsultant)
        .where(RecruiterConsultant.consultant_id == consultant_id)
        .values(is_active=False)
    )

    # Activate or create specified ones
    for rid in recruiter_ids:
        existing_q = await db.execute(
            select(RecruiterConsultant).where(
                RecruiterConsultant.recruiter_id == rid,
                RecruiterConsultant.consultant_id == consultant_id,
            )
        )
        existing = existing_q.scalars().first()
        if existing:
            existing.is_active = True
        else:
            db.add(RecruiterConsultant(
                recruiter_id=rid,
                consultant_id=consultant_id,
                is_active=True,
            ))

    await db.commit()

    # Return the fresh assignment list along with the message so the
    # frontend can update immediately instead of waiting on a refetch
    # (and doesn't need to guess at names from raw ids).
    result = await db.execute(
        select(RecruiterConsultant, User)
        .join(User, User.id == RecruiterConsultant.recruiter_id)
        .where(
            RecruiterConsultant.consultant_id == consultant_id,
            RecruiterConsultant.is_active == True,
        )
        .order_by(User.full_name.asc())
    )
    assigned_recruiters = [
        {"id": str(u.id), "name": u.full_name, "email": u.email} for _, u in result.all()
    ]

    consultant_label = consultant.full_name or consultant.email or f"consultant {consultant_id}"
    return {
        "message": f"Updated recruiter assignments for {consultant_label}.",
        "assigned_recruiters": assigned_recruiters,
    }


class UpdateResumeRichTextPayload(BaseModel):
    resumeRichText: Optional[str] = None

@router.patch("/api/consultants/{consultant_id}/resume-rich-text", response_model=ProfileResponse)
async def update_consultant_resume_rich_text(
    consultant_id: int,
    payload: UpdateResumeRichTextPayload,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    _require_role(current_user, "RECRUITER", "ADMIN")
    if current_user.role == "RECRUITER":
        await _assert_recruiter_mapped(db, current_user.id, consultant_id)

    consultant = await _get_consultant_or_404(db, consultant_id)

    consultant.resume_rich_text = payload.resumeRichText
    await db.commit()
    await db.refresh(consultant)

    count_result = await db.execute(
        select(func.count()).where(ConsultantExperience.consultant_id == consultant_id)
    )
    exp_count = count_result.scalar_one()
    return await _consultant_to_profile_response(db, consultant, exp_count, include_resume_size=False)